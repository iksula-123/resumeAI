"""
DOCX design preservation — Mode 1 ("Preserve Original Format").

The core idea: python-docx lets you read/write a paragraph *run*'s `.text`
without touching its `.font` (family/size/bold/color/etc). So instead of
rebuilding a document from structured content (which is what the
SahiCareer template generator in routers/export.py does — a DESIGN
decision, Mode 2), we open the user's ORIGINAL .docx and only replace the
text of the specific runs that changed, leaving every style, table,
header/footer, margin, and section untouched.

Content layer (what changed) vs design layer (how it looks) are kept
strictly separate: this module never touches fonts, colors, spacing, or
document structure — only `run.text`.
"""
import io

from rapidfuzz import fuzz

_MATCH_THRESHOLD = 80  # fuzzy match confidence for tying an AI bullet back to its source paragraph


def _iter_body_paragraphs(doc):
    """Paragraphs in the main body AND inside table cells, in document order."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def extract_font_color_metadata(data: bytes) -> dict:
    """Best-effort snapshot of the fonts/colors actually used in the document —
    stored on the Resume row (font_metadata/color_metadata) so the app can show
    what will be preserved, and warn if a chosen replacement font is unavailable.
    Never used to CHANGE anything — extraction only."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    fonts, colors, sizes = set(), set(), set()
    for p in _iter_body_paragraphs(doc):
        for run in p.runs:
            if not run.text.strip():
                continue
            if run.font.name:
                fonts.add(run.font.name)
            if run.font.size:
                sizes.add(run.font.size.pt)
            if run.font.color and run.font.color.rgb:
                colors.add(str(run.font.color.rgb))
    return {
        "fonts": sorted(fonts),
        "sizes": sorted(sizes),
        "colors": sorted(colors),
    }


def _paragraph_structure_signature(data: bytes) -> dict:
    """Cheap structural fingerprint used by the preview-validation check —
    NOT a visual diff, just enough to catch "the edit accidentally added/
    removed paragraphs or tables" (the one way this module could change layout)."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    return {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
    }


def _set_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a paragraph's visible text while preserving formatting.

    Puts the new text in the first run (keeping ITS font/size/color/bold —
    representative of the line) and blanks any additional runs rather than
    deleting them, since removing runs can disturb spacing/kerning artifacts
    some documents rely on.
    """
    runs = paragraph.runs
    if not runs:
        # Paragraph had no runs (rare — e.g. an empty placeholder) — add one
        # inheriting the paragraph's style rather than inventing formatting.
        paragraph.add_run(new_text)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _find_best_match(target_text: str, candidates: list[tuple[int, str]]) -> tuple[int, float] | None:
    """Best fuzzy match for `target_text` among (index, text) candidates.
    Returns (index, score) or None if nothing clears the confidence bar.

    Uses fuzz.ratio (overall Levenshtein-based similarity), NOT
    token_set_ratio: token_set_ratio scores a short string as a near-perfect
    match against ANY longer text that merely contains its words as a subset
    — which is exactly wrong here. The AI's parsed "summary"/bullet text can,
    on a bad parse, end up as one giant blob containing the whole resume; with
    token_set_ratio that blob then scores ~100 against EVERY paragraph
    (including the name/heading), and ties resolve to whichever paragraph
    happens to come first — silently overwriting the wrong line. fuzz.ratio
    penalizes length mismatches, so a short paragraph is never mistaken for an
    unrelated blob just because it's textually contained in it. A length-ratio
    guard adds a second, independent line of defense against the same failure
    mode. We want near-identical-text matching here (the AI's own parse
    against its own source), not semantic/subset matching.
    """
    target = (target_text or "").strip()
    if not target:
        return None
    best_idx, best_score = None, 0.0
    for idx, text in candidates:
        candidate = (text or "").strip()
        if not candidate:
            continue
        # Guard against degenerate matches (e.g. a mis-parsed field that's
        # really the whole document) before even scoring: a genuine match
        # shouldn't differ from the candidate paragraph by more than ~3x
        # in length in either direction.
        len_ratio = len(candidate) / len(target) if target else 0
        if len_ratio < 0.3 or len_ratio > 3.0:
            continue
        score = fuzz.ratio(target, candidate)
        if score > best_score:
            best_idx, best_score = idx, score
    if best_idx is not None and best_score >= _MATCH_THRESHOLD:
        return best_idx, best_score
    return None


def apply_content_edits(
    original_bytes: bytes,
    original_content: dict,
    enhanced_content: dict,
) -> tuple[bytes, dict]:
    """Apply AI Improve's content changes onto the ORIGINAL .docx in place.

    Only text that genuinely changed (summary, experience bullets) is
    rewritten, and only inside the run(s) that already held that text —
    every other paragraph, table, section, header/footer, and every font/
    color/size/spacing property is left completely untouched.

    Returns (new_docx_bytes, report) where `report` lists what was matched/
    skipped, for the preview-validation step ("design preserved" check).
    """
    from docx import Document

    doc = Document(io.BytesIO(original_bytes))
    paragraphs = list(_iter_body_paragraphs(doc))
    candidates = [(i, p.text) for i, p in enumerate(paragraphs)]

    report = {"replaced": [], "unmatched": []}

    def replace_if_changed(original_text: str, new_text: str, label: str) -> None:
        original_text = (original_text or "").strip()
        new_text = (new_text or "").strip()
        if not original_text or not new_text or original_text == new_text:
            return
        match = _find_best_match(original_text, candidates)
        if match is None:
            report["unmatched"].append({"label": label, "text": original_text[:80]})
            return
        idx, score = match
        _set_paragraph_text(paragraphs[idx], new_text)
        report["replaced"].append({"label": label, "confidence": round(score)})

    # Summary
    replace_if_changed(original_content.get("summary", ""), enhanced_content.get("summary", ""), "summary")

    # Experience bullets — matched position-for-position within each role, since
    # AI Improve rewrites existing bullets rather than adding/removing them.
    orig_exp = original_content.get("experience") or []
    new_exp = enhanced_content.get("experience") or []
    for oe, ne in zip(orig_exp, new_exp):
        obullets = oe.get("bullets") or []
        nbullets = ne.get("bullets") or []
        for ob, nb in zip(obullets, nbullets):
            replace_if_changed(ob, nb, f"bullet:{(oe.get('position') or '')[:30]}")

    buf = io.BytesIO()
    doc.save(buf)
    new_bytes = buf.getvalue()

    report["structure_before"] = _paragraph_structure_signature(original_bytes)
    report["structure_after"] = _paragraph_structure_signature(new_bytes)
    report["design_preserved"] = (
        report["structure_before"] == report["structure_after"] and not report["unmatched"]
    )
    return new_bytes, report
