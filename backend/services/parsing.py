"""Extract plain text from an uploaded resume (PDF / DOCX / TXT)."""
import io
import re
import zipfile


class InvalidUpload(ValueError):
    """Raised when an uploaded file is not a genuine, readable PDF/DOCX/TXT.

    The message is user-facing and safe to surface directly in an API error.
    """


def validate_upload(filename: str, data: bytes) -> str:
    """Validate that `data` really is the file type its extension claims.

    Guards against files merely *renamed* to .pdf/.docx and against corrupt
    uploads, so we reject them with a clear message instead of failing later
    with a vague "couldn't read text" or sending garbage to the AI.

    Returns the detected extension (".pdf" | ".docx" | ".txt").
    Raises InvalidUpload with a friendly message on any problem.
    """
    name = (filename or "").lower()
    if not data:
        raise InvalidUpload("The file appears to be empty.")

    if name.endswith(".pdf"):
        # Real PDFs start with "%PDF-" (allow a few leading bytes/BOM some tools add).
        if b"%PDF-" not in data[:1024]:
            raise InvalidUpload("This file isn't a valid PDF — it may be renamed or corrupted. "
                                "Please upload a genuine PDF, DOCX, or TXT.")
        return ".pdf"

    if name.endswith(".docx"):
        # DOCX is a ZIP container; verify the ZIP signature and the Word part.
        if not data[:4] == b"PK\x03\x04":
            raise InvalidUpload("This file isn't a valid Word .docx — it may be renamed or corrupted. "
                                "Please upload a genuine PDF, DOCX, or TXT.")
        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                names = set(zf.namelist())
            if "word/document.xml" not in names:
                raise InvalidUpload("This .docx is missing its Word content — it may be an .doc or a "
                                    "different Office file. Please re-save as .docx and try again.")
        except zipfile.BadZipFile:
            raise InvalidUpload("This .docx file is corrupted and can't be opened. "
                                "Please re-save it and try again.")
        return ".docx"

    if name.endswith(".txt"):
        # Reject binary content masquerading as text (NUL bytes never appear in real text).
        if b"\x00" in data[:4096]:
            raise InvalidUpload("This .txt file looks like binary data, not text. "
                                "Please upload a genuine PDF, DOCX, or TXT.")
        return ".txt"

    raise InvalidUpload("Please upload a PDF, DOCX, or TXT file.")


def extract_text(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    text = ""
    try:
        if name.endswith(".pdf"):
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        elif name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(data))
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" ".join(c.text for c in row.cells))
            text = "\n".join(parts)
        else:
            text = data.decode("utf-8", "ignore")
    except Exception:
        # last-ditch: try utf-8 decode
        try:
            text = data.decode("utf-8", "ignore")
        except Exception:
            text = ""
    # normalize whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def quick_contact(text: str) -> dict:
    """Regex fallback to pull obvious contact fields if the AI parse fails."""
    email = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    phone = re.search(r"(\+?\d[\d\s().-]{7,}\d)", text)
    first_line = next((l.strip() for l in text.splitlines() if l.strip()), "")
    return {
        "fullName": first_line[:60] if len(first_line) < 60 else "",
        "email": email.group(0) if email else "",
        "phone": phone.group(0).strip() if phone else "",
        "jobTitle": "", "location": "", "linkedin": "", "github": "", "website": "",
    }


# ── No-AI structured fallback ────────────────────────────────────────────────
# Used when the AI parse call fails/is unavailable (rate limit, no credits,
# network). Splitting into real sections is a MUCH safer degrade than dumping
# the whole resume into "summary" and leaving experience/education/skills
# empty — a candidate's actual experience shouldn't disappear just because the
# AI provider is temporarily down.
_SECTION_HEADERS = {
    "summary": ("summary", "objective", "profile", "about", "about me"),
    "experience": ("experience", "work experience", "employment", "employment history",
                    "professional experience", "work history"),
    "education": ("education", "academic background", "qualifications", "academic qualifications"),
    "skills": ("skills", "technical skills", "core competencies", "key skills", "skill set"),
}
_ALL_HEADERS = {h for headers in _SECTION_HEADERS.values() for h in headers}


def _detect_section(line: str) -> str | None:
    normalized = line.strip().lower().rstrip(":").strip()
    if len(normalized) > 40 or not normalized:
        return None
    for section, headers in _SECTION_HEADERS.items():
        if normalized in headers:
            return section
    return None


def heuristic_structured_parse(text: str) -> dict:
    """No-AI structured parse: detects standard section headers (Experience,
    Education, Skills, Summary) and extracts entries with simple regex
    heuristics. Not as accurate as the AI parser, but preserves the
    candidate's real content structure instead of losing it."""
    lines = [l.rstrip() for l in (text or "").splitlines()]

    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in lines:
        section = _detect_section(line)
        if section:
            current = section
            sections.setdefault(section, [])
            continue
        if current and line.strip():
            sections.setdefault(current, []).append(line.strip())

    personal = quick_contact(text)
    summary = " ".join(sections.get("summary", [])).strip()[:600]

    # Experience: a line that looks like a role/company header (has a
    # separator, "at", or a year) starts a new entry; everything else under it
    # becomes a bullet for that entry.
    experience: list[dict] = []
    current_entry: dict | None = None
    for line in sections.get("experience", []):
        is_bullet = line.startswith(("-", "•", "*")) or bool(re.match(r"^\d+[.)]\s", line))
        looks_like_header = not is_bullet and (
            re.search(r"\s[-–—]\s", line) or re.search(r"\bat\b", line, re.IGNORECASE)
            or re.search(r"\b(19|20)\d{2}\b", line)
        )
        if current_entry is None or looks_like_header:
            current_entry = {"position": "", "company": "", "startDate": "", "endDate": "", "current": False, "bullets": []}
            parts = re.split(r"\s[-–—]\s|,\s+", line, maxsplit=1)
            if len(parts) == 2:
                current_entry["position"], current_entry["company"] = parts[0].strip(), parts[1].strip()
            else:
                current_entry["position"] = line.strip()
            experience.append(current_entry)
        else:
            bullet = re.sub(r"^[-•*]\s*|^\d+[.)]\s*", "", line).strip()
            if bullet:
                current_entry["bullets"].append(bullet)

    # Education: each non-bullet line is treated as its own entry (most
    # resumes list one degree/institution per line).
    education = [
        {"degree": line, "field": "", "institution": "", "startDate": "", "endDate": "", "gpa": ""}
        for line in sections.get("education", [])
        if not line.startswith(("-", "•", "*"))
    ]

    skills: list[str] = []
    for line in sections.get("skills", []):
        for part in re.split(r"[,•|/]", line):
            name = part.strip().strip("-").strip()
            if name and len(name) <= 40:
                skills.append(name)

    return {
        "personalInfo": personal,
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills,
        "projects": [],
        "certifications": [],
        "languages": [],
    }
