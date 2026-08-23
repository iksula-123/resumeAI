"""
Regression tests for docs/ATS_CASE_STUDY_JAVA_DOCX_PARSING_FIX.md (Phase F3).

**Context:** a user uploaded a real DOCX resume
("Java_Developer_Fresher_Resume.docx", found on disk and inspected
directly — see the case study doc for the full trace) against a "frontend
developer" JD and got Experience = 0, 0 bullets detected everywhere, and
achievements bleeding into certifications. Three distinct, confirmed bugs:

1. "INTERNSHIP EXPERIENCE" wasn't a recognized heading alias for the
   "experience" canonical section.
2. Word's NATIVE bulleted/numbered list paragraphs (applied via the
   paragraph's `<w:numPr>` numbering properties, not literal text) were
   invisible to `_is_bullet_line()` — `python-docx`'s `paragraph.text`
   never includes the bullet glyph for these.
3. "ACHIEVEMENTS & EXTRACURRICULAR" wasn't a recognized heading (its "&"
   defeated the unrecognized-heading section-boundary detector), so its
   content bled into the preceding CERTIFICATIONS section.

**Fixture strategy:** rather than depending on the real file (which only
exists on one machine, outside this repo), a DOCX is built PROGRAMMATICALLY
in-memory with `python-docx`, replicating the real file's structurally
relevant shape exactly — including using `add_paragraph(..., style="List
Bullet")` so the fixture has REAL Word-native numPr list formatting, not a
simulated/hand-typed bullet character. This is the same mechanism verified
directly against the real file during this investigation.

No production code beyond `services/ats_engine/resume_parser.py` is
changed by this fix. `section_recognizer.py`, `scoring.py`, `ats_config.py`
are untouched — confirmed via `git diff` in the accompanying report.
"""
import asyncio
import io

import docx
import pytest

from services.ats_engine import resume_parser
from services.ats_engine.resume_parser import extract_text, _is_docx_list_paragraph


def _sync(coro):
    """Deliberately NOT `asyncio.run()` — that helper calls
    `asyncio.set_event_loop(None)` on exit, which clobbers the "current
    event loop" other test files in the same session rely on (this repo's
    conftest.py provides a session-scoped `event_loop` fixture used by
    DB-backed async tests elsewhere — running enough bare `asyncio.run()`
    calls before those tests corrupts that fixture's state and produces
    spurious "no current event loop" failures with no connection to this
    file's own logic). A private loop that's never installed as the
    process-global "current" loop avoids touching that shared state at all."""
    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()


def _build_java_fresher_docx() -> bytes:
    """Programmatic fixture replicating the real resume's structurally
    relevant shape: headings that exercise all three bugs, and REAL
    Word-native bulleted lists (not literal "-"/"•" characters) for every
    bullet — built via python-docx's own "List Bullet" style, which is
    backed by real `<w:numPr>` numbering in the saved file, exactly like
    the real uploaded resume."""
    doc = docx.Document()
    doc.add_paragraph("YOUR NAME")
    doc.add_paragraph("Java Developer | Fresher (B.Tech)")

    doc.add_paragraph("CAREER OBJECTIVE")
    doc.add_paragraph("Motivated B.Tech graduate with a strong foundation in Java and Spring Boot.")

    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("Bachelor of Technology (B.Tech) - Computer Science")
    doc.add_paragraph("Test College, City, State")

    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Languages: Java, SQL, JavaScript, C/C++")
    doc.add_paragraph("Frameworks: Spring Boot, Spring MVC, Hibernate")

    doc.add_paragraph("ACADEMIC PROJECTS")
    doc.add_paragraph("Online Student Management System\tJava, Spring Boot, MySQL")
    doc.add_paragraph("Built a full-stack CRUD application using Spring Boot and MySQL.", style="List Bullet")
    doc.add_paragraph("Designed REST APIs for student and course modules.", style="List Bullet")

    # Bug 1 target — the exact real-world heading that wasn't recognized.
    doc.add_paragraph("INTERNSHIP EXPERIENCE")
    doc.add_paragraph("Java Development Intern - Company Name, City\tJun 2024 - Aug 2024")
    doc.add_paragraph("Assisted in developing and testing REST APIs using Spring Boot.", style="List Bullet")
    doc.add_paragraph("Fixed bugs and wrote unit tests, improving code coverage by 15%.", style="List Bullet")
    doc.add_paragraph("Collaborated with a team of 4 developers using Git.", style="List Bullet")

    doc.add_paragraph("CERTIFICATIONS")
    doc.add_paragraph("Java Programming Masterclass - Udemy (2024)", style="List Bullet")
    doc.add_paragraph("Oracle Certified Foundations Associate: Java - Oracle (2024)", style="List Bullet")

    # Bug 3 target — the exact real-world heading with "&" that bled into
    # CERTIFICATIONS above.
    doc.add_paragraph("ACHIEVEMENTS & EXTRACURRICULAR")
    doc.add_paragraph("Solved 300+ problems on LeetCode covering arrays and trees.", style="List Bullet")
    doc.add_paragraph("Secured 2nd place in a college-level coding hackathon.", style="List Bullet")

    doc.add_paragraph("DECLARATION")
    doc.add_paragraph("I hereby declare that the information provided above is true.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


JAVA_DOCX_BYTES = _build_java_fresher_docx()


def _force_ai_unavailable(monkeypatch):
    async def _none(*args, **kwargs):
        return None
    monkeypatch.setattr(resume_parser, "chat_json", _none)


# ── 1/2/3/4 — Internship Experience is recognized, experience has an
# entry, dates preserved, bullets extracted ─────────────────────────────

def test_internship_experience_heading_is_recognized(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert "experience" in resume["heuristic_sections_found"]


def test_experience_contains_at_least_one_entry(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert len(resume["experience"]) >= 1


def test_internship_dates_are_preserved(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    entry = resume["experience"][0]
    # The pre-existing, separately-documented "title+date on one line"
    # limitation (docs/ATS_CASE_STUDY_POOJA_FORMATTING_FIX.md /
    # ATS_PHASE_F2_FULL_JD_VALIDATION.md §17) means the date range lands in
    # `duration` verbatim rather than being cleanly split from the title —
    # but the dates themselves are real, preserved text, not lost.
    assert "Jun 2024" in entry["duration"] and "Aug 2024" in entry["duration"]


def test_internship_bullets_are_extracted(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    entry = resume["experience"][0]
    assert len(entry["bullets"]) == 3
    assert any("Spring Boot" in b for b in entry["bullets"])
    assert any("4 developers" in b for b in entry["bullets"])


# ── 5/6/7 — native Word bullets detected, project bullets detected, plain
# paragraphs not falsely counted as bullets ──────────────────────────────

def test_native_word_bullets_are_detected_via_numpr():
    """Direct unit check on the detector itself — build one list paragraph
    and one plain paragraph in the same document, confirm only the real
    list item is flagged."""
    doc = docx.Document()
    doc.add_paragraph("A normal heading-like paragraph")
    bullet_p = doc.add_paragraph("A real bulleted item", style="List Bullet")
    plain_p = doc.add_paragraph("A plain paragraph that just happens to be short")
    assert _is_docx_list_paragraph(bullet_p) is True
    assert _is_docx_list_paragraph(plain_p) is False


def test_project_bullets_are_detected(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert len(resume["projects"]) >= 1
    # Bug 2's fix on project grouping: with bullets now detected, the
    # project's real bullet lines fold into ONE project's description
    # instead of each becoming its own fake project.
    combined = " ".join(p["description"] for p in resume["projects"])
    assert "CRUD application" in combined


def test_normal_paragraphs_are_not_incorrectly_counted_as_bullets():
    """Plain, non-list paragraphs (headings, single-line entries) must
    never be treated as bullets just because they sit inside a bulleted
    section — confirmed via the real extraction: only the 7 real
    "List Bullet"-style paragraphs in the fixture produce a leading bullet
    marker, nothing else does."""
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    bullet_marked_lines = [l for l in text.splitlines() if l.startswith("•")]
    assert len(bullet_marked_lines) == 9  # exactly the 9 "List Bullet" paragraphs in the fixture
    assert "YOUR NAME" not in bullet_marked_lines[0]
    assert not any(l.startswith("•") and "INTERNSHIP EXPERIENCE" in l for l in text.splitlines())


# ── 8/9/10 — achievements recognized, not under certifications,
# certifications contain only certification content ────────────────────

def test_achievements_section_is_recognized(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert "achievements" in resume["heuristic_sections_found"]
    assert len(resume["achievements"]) == 2
    assert any("LeetCode" in a for a in resume["achievements"])


def test_achievements_do_not_appear_under_certifications(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    for cert in resume["certifications"]:
        assert "LeetCode" not in cert
        assert "hackathon" not in cert
        assert "ACHIEVEMENTS" not in cert.upper() or "EXTRACURRICULAR" not in cert.upper()


def test_certifications_contain_only_certification_content(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert len(resume["certifications"]) == 2
    assert all("Udemy" in c or "Oracle" in c for c in resume["certifications"])


# ── 11/12 — skills and education still correctly extracted (no
# regression from the heading-alias additions) ──────────────────────────

def test_skills_still_correctly_extracted(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert len(resume["skills"]) > 0
    assert any("Java" in s for s in resume["skills"])


def test_education_remains_correctly_extracted(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert len(resume["education"]) == 1
    assert "B.Tech" in resume["education"][0]["degree"]


# ── 13 — projects grouped correctly (not one fake project per sentence) ──

def test_projects_are_grouped_correctly(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    # KNOWN GAP fixed: before Bug 2's fix, each bullet sentence became its
    # own separate "project" (no bullets were ever detected to group
    # under one project name). Now there's exactly 1 real project with
    # both its bullets folded into one description.
    assert len(resume["projects"]) == 1
    assert "CRUD application" in resume["projects"][0]["description"]
    assert "REST APIs" in resume["projects"][0]["description"]


# ── 14 — no fabricated information ───────────────────────────────────────

def test_no_fabricated_information_is_introduced(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    raw_lower = resume["raw_text"].lower()
    for skill in resume["skills"]:
        assert skill.lower() in raw_lower
    for entry in resume["experience"]:
        for bullet in entry["bullets"]:
            assert bullet.lower() in raw_lower
    for cert in resume["certifications"]:
        assert cert.lower() in raw_lower
    for ach in resume["achievements"]:
        assert ach.lower() in raw_lower


# ── 15 — parser works when AI parsing is unavailable ─────────────────────

def test_parser_works_when_ai_parsing_is_unavailable(monkeypatch):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    assert resume["parsed_by"] == "heuristic"
    assert resume["experience"] and resume["skills"] and resume["education"] and resume["certifications"]


# ── Before/after diagnostic (not a pass/fail assertion — printed for the
# record, matching the fix report's before/after table) ─────────────────

def test_before_after_diagnostic_summary(monkeypatch, capsys):
    _force_ai_unavailable(monkeypatch)
    text = extract_text("java_fresher.docx", JAVA_DOCX_BYTES)
    resume = _sync(resume_parser.parse(text))
    summary = {
        "experience_entries": len(resume["experience"]),
        "experience_bullets": sum(len(e["bullets"]) for e in resume["experience"]),
        "certifications": len(resume["certifications"]),
        "achievements": len(resume["achievements"]),
        "projects": len(resume["projects"]),
        "top_level_bullets": len(resume["bullets"]),
    }
    print("\nAFTER (Phase F3 fix):", summary)
    assert summary["experience_entries"] == 1
    assert summary["experience_bullets"] == 3
    assert summary["certifications"] == 2
    assert summary["achievements"] == 2
    assert summary["projects"] == 1
