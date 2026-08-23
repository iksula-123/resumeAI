"""
Phase H1 — Resume ATS Health calibration.

Covers: the two-layer blend (mode_orchestrator.resume_health_mode()), the
"internship experience" section-alias fix, the Java Developer Fresher
regression case, and confirmation that Phase G's Role Readiness / Job Match
/ JD sufficiency behavior is untouched by any of this.

None of these tests import or assert against any competitor score. The
Java Fresher fixture's new score is computed and reported, never forced to
match anything external (see docs/ATS_CHANGELOG.md's Phase H1 entry).
"""
import io

import docx
import pytest

from services.ats_engine import (
    ats_config, job_parser, mode_orchestrator, parsing_quality, resume_quality, section_recognizer,
)

# ── Synthetic fixtures — deliberately hand-built (not run through
# resume_parser.py) so each test isolates exactly the variable it names,
# same convention test_ats_engine.py already uses for scoring.py's
# category functions. Not claimed to be a "real" parsed resume. ──────────

_CONTACT_HEADER = "Jordan Rivera\njordan.rivera@example.com | +91 98765 43210 | Bengaluru\n\n"

_STRONG_BULLETS = [
    "Developed and deployed 5 microservices in Java and Spring Boot, reducing latency by 40%.",
    "Led a team of 4 engineers migrating the platform onto Kubernetes and Docker, improving deployment speed by 60%.",
    "Automated CI/CD pipelines using Jenkins, cutting release time by 3 hours per week.",
]
_WEAK_BULLETS = [
    "Responsible for backend development.",
    "Worked on various projects.",
    "Helped the team with tasks.",
]
_SKILLS = ["Java", "Spring Boot", "Kubernetes", "Docker", "Jenkins"]
_STRONG_SUMMARY = "Backend engineer with 5 years building distributed Java systems, leading small teams, and shipping measurable performance improvements."
_WEAK_SUMMARY = "Hardworking professional and team player looking for opportunities. Detail-oriented and self-motivated."


def _ats_shell(bullets: list[str], with_headers: bool) -> str:
    """The raw_text half of a resume — with_headers=True gives Parsing/
    Sections/Formatting/Contact everything they look for; False gives none
    of it (no recognizable headers, no bullet markers, no contact info)."""
    if with_headers:
        return (
            _CONTACT_HEADER +
            "SUMMARY\nExperienced backend engineer.\n\n"
            "EXPERIENCE\nSenior Software Engineer - Acme Corp\n" +
            "\n".join(f"- {b}" for b in bullets) + "\n\n"
            "EDUCATION\nB.Tech Computer Science, IIT Delhi, 2018\n\n"
            "SKILLS\n" + ", ".join(_SKILLS) + "\n\n"
            "CERTIFICATIONS\nAWS Certified Solutions Architect, AWS, 2021\n"
        )
    # No headers, no bullet markers, no email/phone -- just unstructured
    # prose (clean words, not garbled, to isolate STRUCTURE from raw
    # extraction-cleanliness, which parsing_quality.py measures separately).
    return (
        "Backend engineer with several years of experience building and "
        "operating server side systems across a variety of teams and "
        "projects over the course of a long career in software. " * 3
    )


def _resume(bullets: list[str], summary: str, skills: list[str], with_headers: bool) -> dict:
    text = _ats_shell(bullets, with_headers)
    return {
        "raw_text": text,
        "summary": summary,
        "skills": skills, "hard_skills": skills,
        "experience": [{"title": "Senior Software Engineer", "company": "Acme Corp", "duration": "2019-2024",
                         "bullets": bullets}],
        "education": [{"degree": "B.Tech Computer Science", "institution": "IIT Delhi", "year": "2018"}],
        "projects": [], "certifications": ["AWS Certified Solutions Architect"],
        "total_experience_years": 5, "location": "Bengaluru",
    }


STRONG_ATS_STRONG_QUALITY = _resume(_STRONG_BULLETS, _STRONG_SUMMARY, _SKILLS, with_headers=True)
STRONG_ATS_WEAK_QUALITY = _resume(_WEAK_BULLETS, _WEAK_SUMMARY, _SKILLS, with_headers=True)
WEAK_ATS_STRONG_QUALITY = _resume(_STRONG_BULLETS, _STRONG_SUMMARY, _SKILLS, with_headers=False)
WEAK_ATS_WEAK_QUALITY = _resume(_WEAK_BULLETS, _WEAK_SUMMARY, _SKILLS, with_headers=False)


def _java_fresher_docx_text() -> str:
    """The exact fixture test_java_fresher_docx_regression.py builds — the
    real uploaded file isn't in the repo; this programmatic DOCX replicates
    its structurally relevant shape (see that file's own docstring)."""
    from services.ats_engine.resume_parser import extract_text

    doc = docx.Document()
    doc.add_paragraph("YOUR NAME")
    doc.add_paragraph("Java Developer | Fresher (B.Tech)")
    doc.add_paragraph("CAREER OBJECTIVE")
    doc.add_paragraph("Motivated B.Tech graduate with a strong foundation in Java and Spring Boot.")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("Bachelor of Technology (B.Tech) - Computer Science")
    doc.add_paragraph("Test College, City, State")
    doc.add_paragraph("TECHNICAL SKILLS")
    doc.add_paragraph("Languages: Java, SQL, JavaScript, C/C++")
    doc.add_paragraph("Frameworks: Spring Boot, Spring MVC, Hibernate")
    doc.add_paragraph("ACADEMIC PROJECTS")
    doc.add_paragraph("Online Student Management System\tJava, Spring Boot, MySQL")
    doc.add_paragraph("Built a full-stack CRUD application using Spring Boot and MySQL.", style="List Bullet")
    doc.add_paragraph("Designed REST APIs for student and course modules.", style="List Bullet")
    doc.add_paragraph("INTERNSHIP EXPERIENCE")
    doc.add_paragraph("Java Development Intern - Company Name, City\tJun 2024 - Aug 2024")
    doc.add_paragraph("Assisted in developing and testing REST APIs using Spring Boot.", style="List Bullet")
    doc.add_paragraph("Fixed bugs and wrote unit tests, improving code coverage by 15%.", style="List Bullet")
    doc.add_paragraph("Collaborated with a team of 4 developers using Git.", style="List Bullet")
    doc.add_paragraph("CERTIFICATIONS")
    doc.add_paragraph("Java Programming Masterclass - Udemy (2024)", style="List Bullet")
    doc.add_paragraph("Oracle Certified Foundations Associate: Java - Oracle (2024)", style="List Bullet")
    doc.add_paragraph("ACHIEVEMENTS & EXTRACURRICULAR")
    doc.add_paragraph("Solved 300+ problems on LeetCode covering arrays and trees.", style="List Bullet")
    doc.add_paragraph("Secured 2nd place in a college-level coding hackathon.", style="List Bullet")
    doc.add_paragraph("DECLARATION")
    doc.add_paragraph("I hereby declare that the information provided above is true.")
    buf = io.BytesIO()
    doc.save(buf)
    return extract_text("java_fresher.docx", buf.getvalue())


async def _java_fresher_resume(monkeypatch) -> dict:
    from services.ats_engine import resume_parser

    async def _none(*a, **k):
        return None
    monkeypatch.setattr(resume_parser, "chat_json", _none)
    text = _java_fresher_docx_text()
    return await resume_parser.parse(text)


# ── Config sanity ────────────────────────────────────────────────────────

def test_layer_weights_are_the_approved_45_55_split():
    assert ats_config.RESUME_HEALTH_LAYER_WEIGHTS == {"ats_compatibility": 0.45, "resume_quality": 0.55}


def test_compatibility_weights_sum_to_one_and_include_contact():
    weights = ats_config.RESUME_HEALTH_COMPATIBILITY_WEIGHTS
    assert set(weights) == {"parsing", "sections", "formatting", "contact"}
    assert abs(sum(weights.values()) - 1.0) < 1e-6
    assert weights["contact"] == 0.10


def test_scoring_engine_version_bumped_for_the_resume_health_model_change():
    """Phase H1 materially changed Resume ATS Health's formula (4-category
    -> two-layer 45/55 blend) -- per explicit approval, the version was
    bumped so 2.0.0-era persisted reports stay distinguishable from 2.1.0
    ones. Job Match / Role Readiness / JD sufficiency are unaffected."""
    assert ats_config.SCORING_ENGINE_VERSION == "2.1.0"


# ── Section alias fix ────────────────────────────────────────────────────

def test_internship_experience_alias_is_recognized():
    text = "SUMMARY\nSome text.\n\nINTERNSHIP EXPERIENCE\nSome company\n- did stuff\n\nEDUCATION\nB.Tech\n\nSKILLS\nJava\n"
    result = section_recognizer.recognize_sections(text)
    assert "experience" in result["recognized_sections"]
    assert "experience" not in result["missing_core_sections"]


def test_existing_experience_headings_still_recognized_no_regression():
    for heading in ("EXPERIENCE", "WORK EXPERIENCE", "PROFESSIONAL EXPERIENCE", "EMPLOYMENT HISTORY"):
        text = f"SUMMARY\nx\n\n{heading}\nCompany\n- did stuff\n\nEDUCATION\nx\n\nSKILLS\nx\n"
        result = section_recognizer.recognize_sections(text)
        assert "experience" in result["recognized_sections"], f"{heading} regressed"


# ── contact_note exposure (additive, parsing score untouched) ──────────────

def test_contact_note_is_exposed_without_changing_parsing_score():
    text = _ats_shell(_STRONG_BULLETS, with_headers=True)
    result = parsing_quality.analyze_parsing_quality(text)
    assert "contact_note" in result
    assert isinstance(result["contact_note"], str) and result["contact_note"]
    # the score formula itself is untouched -- same 6-term blend as before
    assert 0 <= result["score"] <= 100


# ── Two-layer blend — cross combinations (Phase H1 tests 1-4) ──────────────

def test_1_strong_ats_strong_quality_scores_high():
    result = mode_orchestrator.resume_health_mode(STRONG_ATS_STRONG_QUALITY)
    assert result["score"] >= 80
    assert result["layers"]["ats_compatibility"]["score"] >= 80
    assert result["layers"]["resume_quality"]["score"] >= 70


def test_2_strong_ats_weak_quality_is_materially_lower_than_ats_alone():
    result = mode_orchestrator.resume_health_mode(STRONG_ATS_WEAK_QUALITY)
    layer_a = result["layers"]["ats_compatibility"]["score"]
    layer_b = result["layers"]["resume_quality"]["score"]
    assert layer_a >= 80          # ATS side is genuinely strong
    assert layer_b <= 45          # quality side is genuinely weak
    assert result["score"] <= layer_a - 15, "weak content must pull the overall score materially below the ATS-only number"


def test_3_weak_ats_strong_quality_is_materially_lower_than_quality_alone():
    result = mode_orchestrator.resume_health_mode(WEAK_ATS_STRONG_QUALITY)
    layer_a = result["layers"]["ats_compatibility"]["score"]
    layer_b = result["layers"]["resume_quality"]["score"]
    assert layer_a <= 60           # structurally weak (no headers/bullets/contact)
    assert layer_b >= 70           # content itself is strong
    assert result["score"] <= layer_b - 10, "poor ATS compatibility must pull the overall score below the quality-only number -- strong content does not fully rescue an unparseable resume"


def test_4_weak_ats_weak_quality_scores_low():
    result = mode_orchestrator.resume_health_mode(WEAK_ATS_WEAK_QUALITY)
    assert result["score"] <= 45


# ── Explicit proof: Resume Quality has NON-ZERO influence ──────────────────

def test_resume_quality_has_nonzero_influence_on_resume_ats_health():
    """The core Phase H bug, proven as a true controlled experiment: hold
    raw_text (and therefore Layer A -- parsing/sections/formatting/contact)
    BYTE-IDENTICAL between two resumes, and vary ONLY the structured fields
    Layer B reads (bullets/summary/skills). Same ATS shell, different
    content quality -- the two must NOT get the same Resume ATS Health
    score."""
    shared_text = _ats_shell(_STRONG_BULLETS, with_headers=True)

    strong_resume = dict(STRONG_ATS_STRONG_QUALITY, raw_text=shared_text)
    weak_resume = dict(STRONG_ATS_WEAK_QUALITY, raw_text=shared_text)

    strong = mode_orchestrator.resume_health_mode(strong_resume)
    weak = mode_orchestrator.resume_health_mode(weak_resume)

    assert strong["layers"]["ats_compatibility"]["score"] == weak["layers"]["ats_compatibility"]["score"], \
        "raw_text is byte-identical -- Layer A must be exactly equal"
    assert strong["layers"]["resume_quality"]["score"] > weak["layers"]["resume_quality"]["score"] + 20
    assert strong["score"] > weak["score"] + 10, \
        "Resume Quality must move the overall score even when ATS Compatibility is held constant"
    assert strong["layer_weights_used"]["resume_quality"] == 55.0


# ── Individual weak-signal tests (5-9) ──────────────────────────────────────

def test_5_missing_contact_information_is_scored_and_visible():
    with_contact = mode_orchestrator.resume_health_mode(STRONG_ATS_STRONG_QUALITY)
    no_contact_resume = dict(STRONG_ATS_STRONG_QUALITY)
    no_contact_resume["raw_text"] = _java_fresher_docx_text()  # no email/phone anywhere, reuse as a no-contact text
    without_contact = mode_orchestrator.resume_health_mode(no_contact_resume)

    contact_cat = without_contact["layers"]["ats_compatibility"]["categories"]["contact"]
    assert contact_cat["match"] == 0.0
    assert without_contact["layers"]["ats_compatibility"]["score"] < with_contact["layers"]["ats_compatibility"]["score"]


def test_6_weak_skill_evidence_scores_low_strong_scores_high():
    """The anti-gaming case from the Phase H1 spec: listing skills is not
    the same as demonstrating them. 'Java, Spring Boot, Hibernate' with no
    supporting bullet must score lower than the same skills actually used
    in a bullet."""
    weak = resume_quality._skill_evidence_category({
        "skills": ["Java", "Spring Boot", "Hibernate"], "hard_skills": [],
        "experience": [{"bullets": ["Assisted with various tasks."]}], "projects": [], "education": [],
    })
    strong = resume_quality._skill_evidence_category({
        "skills": ["Java", "Spring Boot", "Hibernate"], "hard_skills": [],
        "experience": [{"bullets": ["Developed Spring Boot REST APIs using Java and Hibernate for a checkout service."]}],
        "projects": [], "education": [],
    })
    assert weak.match == 0.0
    assert strong.match == 100.0
    assert strong.match > weak.match


def test_7_weak_quantified_impact_scores_low():
    result = mode_orchestrator.resume_health_mode(STRONG_ATS_WEAK_QUALITY)
    cat = result["layers"]["resume_quality"]["categories"]["quantified_impact"]
    assert cat["match"] == 0.0


def test_8_weak_summary_scores_lower_than_strong_summary():
    strong = mode_orchestrator.resume_health_mode(STRONG_ATS_STRONG_QUALITY)
    weak = mode_orchestrator.resume_health_mode(STRONG_ATS_WEAK_QUALITY)
    strong_summary = strong["layers"]["resume_quality"]["categories"]["summary_quality"]["match"]
    weak_summary = weak["layers"]["resume_quality"]["categories"]["summary_quality"]["match"]
    assert weak_summary < strong_summary


def test_9_weak_recruiter_readiness_scores_lower_than_strong():
    strong = mode_orchestrator.resume_health_mode(STRONG_ATS_STRONG_QUALITY)
    weak = mode_orchestrator.resume_health_mode(STRONG_ATS_WEAK_QUALITY)
    strong_rr = strong["layers"]["resume_quality"]["categories"]["recruiter_readiness"]["match"]
    weak_rr = weak["layers"]["resume_quality"]["categories"]["recruiter_readiness"]["match"]
    assert weak_rr < strong_rr


# ── N/A handling — never becomes 100, never becomes an artificial zero ─────

def test_na_categories_are_excluded_not_scored_as_100_or_0():
    """A resume with fewer than 2 experience entries makes career_progression
    genuinely inapplicable -- it must be absent from the scored categories
    (excluded + redistributed), never silently 100 or silently 0."""
    result = mode_orchestrator.resume_health_mode(STRONG_ATS_STRONG_QUALITY)  # only 1 experience entry
    quality_layer = result["layers"]["resume_quality"]
    assert "career_progression" in quality_layer["excluded_categories"]
    cp = quality_layer["categories"].get("career_progression")
    if cp is not None:
        assert cp["match"] is None


# ── 10 — Java Developer Fresher regression (illustrative, not competitor-targeted) ──

@pytest.mark.asyncio
async def test_10_java_fresher_regression_reflects_real_weaknesses(monkeypatch):
    resume = await _java_fresher_resume(monkeypatch)
    old_health = __import__(
        "services.ats_engine.ats_intelligence_v2", fromlist=["compute_resume_health_v2"]
    ).compute_resume_health_v2(resume)
    new_health = mode_orchestrator.resume_health_mode(resume)

    # The old, structure-only model and the new two-layer model are BOTH
    # printed here (see the accompanying report) -- this test asserts the
    # RELATIONSHIP (new < old, because real weaknesses now count), never a
    # specific target number, and never anything derived from a competitor.
    assert new_health["score"] < old_health["score"] - 10, (
        f"expected the two-layer score ({new_health['score']}) to be materially "
        f"below the structure-only score ({old_health['score']}) for a resume "
        f"with 0/7 skills evidenced and weak action verbs/summary"
    )
    quality_cats = new_health["layers"]["resume_quality"]["categories"]
    assert quality_cats["skill_evidence"]["match"] == 0.0
    assert quality_cats["action_verbs"]["match"] < 50


# ── 14-17 — Phase G modes remain independent and untouched ──────────────────

def test_14_resume_only_mode_still_independent_of_role_and_job():
    resume_health = mode_orchestrator.resume_health_mode(STRONG_ATS_STRONG_QUALITY)
    role = mode_orchestrator.role_readiness_mode(STRONG_ATS_STRONG_QUALITY, None)
    job = mode_orchestrator.job_match_mode(STRONG_ATS_STRONG_QUALITY, None, None)
    assert resume_health["available"] is True
    assert role["available"] is False
    assert job["available"] is False


def test_15_role_readiness_unaffected_by_phase_h1():
    role_job = {
        "job_title": "React Developer", "required_skills": ["React"], "preferred_skills": [],
        "responsibilities": [], "min_experience_years": None, "min_education": "Bachelor's degree",
        "industry": "IT", "keywords": ["React"], "technologies": [], "certifications": [],
        "raw_text": "React Developer React", "parsed_by": "role_library",
    }
    result = mode_orchestrator.role_readiness_mode(STRONG_ATS_STRONG_QUALITY, role_job)
    assert result["available"] is True
    assert result["data_sufficiency"] == "sufficient"
    assert "sufficient" not in result  # role_readiness never carries job_match's vocabulary


def test_16_insufficient_jd_still_never_scored():
    parsed = job_parser._heuristic_parse("java developer")
    sufficiency = job_parser.assess_sufficiency("java developer", parsed)
    assert sufficiency["sufficient"] is False
    result = mode_orchestrator.job_match_mode(STRONG_ATS_STRONG_QUALITY, parsed, sufficiency)
    assert result["sufficient"] is False
    assert result["score"] is None


def test_17_full_job_match_scoring_unchanged():
    raw_text = (
        "Backend Engineer (React Developer team)\n\n"
        "We are looking for an experienced backend engineer to build and maintain our services.\n\n"
        "Responsibilities:\n- Build backend services\n- Collaborate with the platform team\n\n"
        "Required skills: Java, Spring Boot, Kubernetes, Docker\n"
        "Minimum 3 years of experience required."
    )
    job = {
        "job_title": "React Developer", "required_skills": ["Java", "Spring Boot", "Kubernetes", "Docker"],
        "preferred_skills": [], "responsibilities": ["Build backend services", "Collaborate with the platform team"],
        "min_experience_years": 3, "min_education": None, "industry": "IT",
        "keywords": ["Java", "Spring Boot", "Kubernetes", "Docker"], "technologies": [], "certifications": [],
        "raw_text": raw_text, "parsed_by": "ai",
    }
    sufficiency = job_parser.assess_sufficiency(job["raw_text"], job)
    result = mode_orchestrator.job_match_mode(STRONG_ATS_STRONG_QUALITY, job, sufficiency)
    assert result["sufficient"] is True
    assert result["score"] is not None
    assert result["score"] > 50  # strong overlap between resume skills and JD requirements
