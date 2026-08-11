"""
Phase 2 — the five new templates (tech-stack, fresher, academic, healthcare,
international). Covers what test_template_registry.py's generic
parametrized suite doesn't: empty-data handling, long-content/multi-page
behavior, Fresher's adaptive section reordering, and Tech Stack's skill
categorization.
"""
import io

import pytest
from docx import Document
from pypdf import PdfReader

from routers.export import TEMPLATE_BUILDERS, TEMPLATE_SPECS, SINGLE_COLUMN_CONFIGS
from services.skill_categories import categorize_skill, group_skills_by_category

NEW_TEMPLATE_IDS = ["tech-stack", "fresher", "academic", "healthcare", "international"]

MINIMAL_CONTENT = {
    "personalInfo": {"fullName": "Alex Doe", "email": "alex@example.com"},
}


# ── Empty data: no crash, no dangling empty sections ────────────────────────

@pytest.mark.parametrize("template_id", NEW_TEMPLATE_IDS)
def test_handles_completely_empty_optional_sections(template_id):
    """Only personalInfo present — every optional section must be silently
    skipped, not rendered as an empty heading."""
    sections = TEMPLATE_SPECS[template_id]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS[template_id].pdf(MINIMAL_CONTENT, "Test", sections)
    assert pdf_bytes[:5] == b"%PDF-"
    text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages).upper()
    assert "ALEX DOE" in text

    labels = SINGLE_COLUMN_CONFIGS[template_id]["labels"]
    for key, heading in labels.items():
        if key == "summary":
            continue  # personalInfo itself isn't a labeled section
        assert heading.upper() not in text, (
            f"[{template_id}] rendered an empty '{heading}' heading with no data behind it"
        )

    docx_bytes = TEMPLATE_BUILDERS[template_id].docx(MINIMAL_CONTENT, "Test", sections)
    assert docx_bytes[:2] == b"PK"
    docx_text = "\n".join(p.text for p in Document(io.BytesIO(docx_bytes)).paragraphs).upper()
    for key, heading in labels.items():
        if key == "summary":
            continue
        assert heading.upper() not in docx_text, f"[{template_id}] DOCX rendered an empty '{heading}' heading"


@pytest.mark.parametrize("template_id", NEW_TEMPLATE_IDS)
def test_no_experience_does_not_crash(template_id):
    """Explicit empty experience list — the case Fresher is built around,
    but every template must handle it without error."""
    content = {**MINIMAL_CONTENT, "experience": []}
    sections = TEMPLATE_SPECS[template_id]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS[template_id].pdf(content, "Test", sections)
    assert pdf_bytes[:5] == b"%PDF-"


# ── Fresher: adaptive section ordering ──────────────────────────────────────

def test_fresher_promotes_education_and_projects_when_no_experience():
    content = {
        "personalInfo": {"fullName": "Student One"},
        "summary": "Aspiring engineer.",
        "education": [{"institution": "State University", "degree": "B.Sc"}],
        "projects": [{"name": "Weather App", "description": "A weather app."}],
    }
    sections = TEMPLATE_SPECS["fresher"]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS["fresher"].pdf(content, "Test", sections)
    text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages).upper()
    assert "INTERNSHIPS & EXPERIENCE" not in text, "empty Experience heading should never render"
    # Education must appear before Projects, and both before any experience heading (there is none)
    assert text.index("EDUCATION") < text.index("PROJECTS")


def test_fresher_shows_experience_when_present():
    content = {
        "personalInfo": {"fullName": "Student Two"},
        "experience": [{"position": "Intern", "company": "Acme", "startDate": "2024", "endDate": "2024",
                          "current": False, "bullets": ["Did an internship task."]}],
        "education": [{"institution": "State University", "degree": "B.Sc"}],
    }
    sections = TEMPLATE_SPECS["fresher"]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS["fresher"].pdf(content, "Test", sections)
    text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages).upper()
    assert "INTERNSHIPS & EXPERIENCE" in text
    assert "DID AN INTERNSHIP TASK" in text


# ── Tech Stack: skill categorization ────────────────────────────────────────

def test_skill_categorization_common_cases():
    assert categorize_skill("React") == "Frontend"
    assert categorize_skill("PostgreSQL") == "Database"
    assert categorize_skill("Docker") == "DevOps"
    assert categorize_skill("AWS") == "Cloud"
    assert categorize_skill("Jest") == "Testing"
    assert categorize_skill("C#") == "Languages"
    assert categorize_skill(".NET") == "Backend"
    assert categorize_skill("Leadership") == "Other"  # unmatched → Other, not miscategorized


def test_skill_categorization_no_false_positive_substrings():
    """Regression guard for the exact bug found during Phase 2 build: short
    keywords like "c" or "sql" must not match as substrings of unrelated
    words ("react" contains "c", "postgresql" contains "sql")."""
    assert categorize_skill("React") != "Languages"
    assert categorize_skill("Docker") != "Languages"
    assert categorize_skill("PostgreSQL") != "Languages"


def test_tech_stack_pdf_groups_skills_by_category():
    content = {
        "personalInfo": {"fullName": "Dev One"},
        "skills": [{"name": "React"}, {"name": "AWS"}, {"name": "PostgreSQL"}],
    }
    sections = TEMPLATE_SPECS["tech-stack"]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS["tech-stack"].pdf(content, "Test", sections)
    text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages)
    assert "Frontend: React" in text
    assert "Cloud: AWS" in text
    assert "Database: PostgreSQL" in text


def test_group_skills_by_category_omits_empty_buckets():
    groups = group_skills_by_category(["React", "AWS"])
    categories = [cat for cat, _ in groups]
    assert "Frontend" in categories and "Cloud" in categories
    assert "Testing" not in categories  # no testing skills supplied


# ── Long content: multi-page, no crash, no clipping-induced failure ────────

def _long_content() -> dict:
    return {
        "personalInfo": {"fullName": "Long Content Candidate", "jobTitle": "Senior Everything Engineer " * 3,
                         "email": "long@example.com", "phone": "+91 90000 00000", "location": "Bengaluru, India"},
        "summary": "Experienced professional. " * 60,
        "experience": [
            {
                "position": f"Software Engineer Level {i} — Extremely Long Job Title For Testing Purposes",
                "company": f"Very Long Company Name Private Limited {i}",
                "startDate": f"{2010 + i}", "endDate": f"{2011 + i}", "current": False,
                "bullets": [f"Delivered feature number {j} with significant measurable impact on the business." for j in range(4)],
            }
            for i in range(12)
        ],
        "education": [{"institution": f"University {i}", "degree": "B.Tech", "field": "CS", "endDate": str(2010 + i)} for i in range(3)],
        "skills": [{"name": f"Skill{i}"} for i in range(25)],
        "projects": [{"name": f"Project {i}", "technologies": "Python, React", "description": "A long project description. " * 5} for i in range(12)],
        "certifications": [{"name": f"Certification {i}", "issuer": "Issuer", "date": "2022"} for i in range(12)],
        "achievements": [f"Achievement number {i} with a lot of detail." for i in range(10)],
        "languages": [{"name": f"Language{i}", "proficiency": "Fluent"} for i in range(6)],
        "interests": [f"Interest{i}" for i in range(10)],
    }


@pytest.mark.parametrize("template_id", NEW_TEMPLATE_IDS)
def test_long_content_does_not_crash_and_paginates(template_id):
    content = _long_content()
    sections = TEMPLATE_SPECS[template_id]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS[template_id].pdf(content, "Test", sections)
    assert pdf_bytes[:5] == b"%PDF-"
    reader = PdfReader(io.BytesIO(pdf_bytes))
    assert len(reader.pages) > 1, f"[{template_id}] 12 experience entries + 12 projects + 12 certs should span multiple pages"

    text = "\n".join((p.extract_text() or "") for p in reader.pages)
    assert "Long Content Candidate" in text
    assert "Skill24" in text or "skill24" in text.lower()  # last skill wasn't clipped

    docx_bytes = TEMPLATE_BUILDERS[template_id].docx(content, "Test", sections)
    assert docx_bytes[:2] == b"PK"
    doc = Document(io.BytesIO(docx_bytes))
    docx_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Project 11" in docx_text  # last project wasn't dropped
    assert "Certification 11" in docx_text  # last certification wasn't dropped
