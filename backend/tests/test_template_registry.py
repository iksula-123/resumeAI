"""
Template-aware export foundation — registry/smoke tests.

Covers all 10 template_ids end-to-end: the shared spec recognizes it, a
builder is registered for it, PDF/DOCX generation doesn't crash, output is
non-empty, and — the actual point of this whole foundation — switching
template_id never silently drops a candidate's projects/certifications/
achievements/languages/interests.

Does NOT test pixel-level visual output (no such tooling exists in this
repo, and Option A — headless-browser rendering — was explicitly ruled out;
see the Phase 1 architecture discussion). Structural/textual verification via
pypdf/python-docx is the ceiling here.
"""
import io

import pytest
from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader

from routers.export import (
    TEMPLATE_SPECS, TEMPLATE_BUILDERS, DEFAULT_TEMPLATE_ID, SINGLE_COLUMN_CONFIGS,
    ExportRequest, _resolve_template_id,
)

ALL_TEMPLATE_IDS = [
    "modern", "professional", "minimal", "creative", "executive",        # classic family
    "tech-stack", "fresher", "academic", "healthcare", "international",  # single-column family (Phase 2)
]


# Each classic template now has its OWN heading labels (mirroring its
# frontend component's actual copy — ModernTemplate.tsx says "Work
# Experience", ExecutiveTemplate.tsx says "Executive Profile"/"Career
# History", CreativeTemplate.tsx says "About Me", etc.) instead of every
# classic template sharing one generic set. Minimal renders the summary as
# an unheaded left-border quote (matching MinimalTemplate.tsx exactly) — no
# "summary" key for it is intentional, not an omission.
_CLASSIC_HEADING_MAPS: dict[str, dict[str, str]] = {
    "modern": {
        "summary": "PROFESSIONAL SUMMARY", "experience": "WORK EXPERIENCE", "education": "EDUCATION",
        "skills": "SKILLS", "projects": "PROJECTS", "certifications": "CERTIFICATIONS",
        "achievements": "ACHIEVEMENTS", "languages": "LANGUAGES", "interests": "INTERESTS",
    },
    "professional": {
        "summary": "PROFESSIONAL SUMMARY", "experience": "EXPERIENCE", "education": "EDUCATION",
        "skills": "SKILLS", "projects": "PROJECTS", "certifications": "CERTIFICATIONS",
        "achievements": "ACHIEVEMENTS", "languages": "LANGUAGES", "interests": "INTERESTS",
    },
    "minimal": {
        "experience": "EXPERIENCE", "education": "EDUCATION",
        "skills": "SKILLS", "projects": "PROJECTS", "certifications": "CERTIFICATIONS",
        "achievements": "ACHIEVEMENTS", "languages": "LANGUAGES", "interests": "INTERESTS",
    },
    "creative": {
        "summary": "ABOUT ME", "experience": "EXPERIENCE", "education": "EDUCATION",
        "skills": "SKILLS", "projects": "PROJECTS", "certifications": "CERTIFICATIONS",
        "achievements": "ACHIEVEMENTS", "languages": "LANGUAGES", "interests": "INTERESTS",
    },
    "executive": {
        "summary": "EXECUTIVE PROFILE", "experience": "CAREER HISTORY", "education": "EDUCATION",
        "skills": "CORE COMPETENCIES", "projects": "PROJECTS", "certifications": "CERTIFICATIONS",
        "achievements": "KEY ACHIEVEMENTS", "languages": "LANGUAGES", "interests": "INTERESTS",
    },
}


def _docx_all_text(doc) -> str:
    """python-docx's `doc.paragraphs` only walks TOP-LEVEL paragraphs — table
    CELL content (used by the sidebar/2-column classic layouts: modern,
    executive, creative) is invisible to it. Table cells can themselves
    contain nested tables (none of the current layouts do, but walking
    recursively costs nothing and avoids a silent gap if one ever does)."""
    parts = [p.text for p in doc.paragraphs]

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
                for nested in cell.tables:
                    walk_table(nested)

    for t in doc.tables:
        walk_table(t)
    return "\n".join(parts)


def _expected_heading_map(template_id: str) -> dict[str, str]:
    """Single-column templates use their own configured labels (e.g.
    fresher's summary section is headed "Career Objective", not "Summary")
    — see SINGLE_COLUMN_CONFIGS in routers/export.py. Classic templates each
    have their own real design-matched labels — see _CLASSIC_HEADING_MAPS
    above. Sourcing the expectation from the same config/labels the builder
    actually uses is what makes this a real section-parity check instead of
    an assumption."""
    if template_id in SINGLE_COLUMN_CONFIGS:
        return {k: v.upper() for k, v in SINGLE_COLUMN_CONFIGS[template_id]["labels"].items()}
    return _CLASSIC_HEADING_MAPS[template_id]

# Every optional section populated, so "does this template drop data" is
# actually exercised for every template_id, not just the ones with rich specs.
SAMPLE_CONTENT = {
    "personalInfo": {
        "fullName": "Priya Sharma", "jobTitle": "Product Manager",
        "email": "priya@example.com", "phone": "+91 90000 00000", "location": "Bengaluru",
    },
    "summary": "Product leader with 8 years shipping consumer apps.",
    "experience": [{
        "position": "Product Manager", "company": "Zenith Corp",
        "startDate": "2019", "endDate": "2024", "current": False,
        "bullets": ["Owned roadmap for a 2M-user consumer app."],
    }],
    "education": [{"institution": "IIM Bengaluru", "degree": "MBA", "field": "", "endDate": "2018"}],
    "skills": [{"name": "Product Strategy"}, {"name": "SQL"}],
    "projects": [{
        "name": "Internal Analytics Dashboard", "technologies": "React, Python",
        "description": "Built a dashboard used by 50+ PMs.",
    }],
    "certifications": [{"name": "PMP", "issuer": "PMI", "date": "2021"}],
    "achievements": ["Grew MAU by 40 percent in one year"],
    "languages": [{"name": "English", "proficiency": "Native"}, {"name": "Hindi", "proficiency": "Fluent"}],
    "interests": ["Chess", "Reading"],
}

# ── 1. Registry completeness ────────────────────────────────────────────────

def test_all_ids_match_the_shared_spec_file():
    """Guards against this test file's ALL_TEMPLATE_IDS and
    shared/template-specs.json silently drifting apart."""
    assert set(TEMPLATE_SPECS.keys()) == set(ALL_TEMPLATE_IDS)


@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_template_is_recognized(template_id):
    assert template_id in TEMPLATE_SPECS, f"'{template_id}' missing from shared/template-specs.json"


@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_spec_has_required_fields(template_id):
    spec = TEMPLATE_SPECS[template_id]
    for field in ("id", "name", "category", "description", "layoutFamily", "accent",
                  "font", "sections", "atsCompatibility", "bestFor", "status"):
        assert field in spec, f"'{template_id}' spec missing required field '{field}'"
    assert spec["id"] == template_id
    assert isinstance(spec["sections"], list) and spec["sections"]
    assert isinstance(spec["bestFor"], list) and spec["bestFor"]
    assert spec["status"] in ("available", "coming_soon")


@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_builder_is_registered(template_id):
    assert template_id in TEMPLATE_BUILDERS, f"'{template_id}' has no registered builder"
    builder = TEMPLATE_BUILDERS[template_id]
    assert callable(builder.pdf) and callable(builder.docx)


# ── 2. PDF/DOCX generation doesn't crash, output is non-empty ──────────────

@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_pdf_generation_does_not_crash(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test Resume", sections)
    assert isinstance(pdf_bytes, (bytes, bytearray))
    assert len(pdf_bytes) > 500
    assert pdf_bytes[:5] == b"%PDF-"


@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_docx_generation_does_not_crash(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    docx_bytes = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test Resume", sections)
    assert isinstance(docx_bytes, (bytes, bytearray))
    assert len(docx_bytes) > 1000
    assert docx_bytes[:2] == b"PK"  # valid zip/docx container


# ── 3. Section parity: what the spec declares is what actually renders ─────

@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_expected_sections_present_in_pdf(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    heading_map = _expected_heading_map(template_id)
    pdf_bytes = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test Resume", sections)
    text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages).upper()
    assert "PRIYA SHARMA" in text
    for key, heading in heading_map.items():
        if key in sections:
            assert heading in text, f"[{template_id}] declared section '{key}' (heading '{heading}') missing from PDF"


@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_expected_sections_present_in_docx(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    heading_map = _expected_heading_map(template_id)
    docx_bytes = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test Resume", sections)
    text = _docx_all_text(Document(io.BytesIO(docx_bytes))).upper()
    assert "PRIYA SHARMA" in text
    for key, heading in heading_map.items():
        if key in sections:
            assert heading in text, f"[{template_id}] declared section '{key}' (heading '{heading}') missing from DOCX"


# ── 4. The actual point of this foundation: no data loss on template switch ─

@pytest.mark.parametrize("template_id", ALL_TEMPLATE_IDS)
def test_no_data_loss_when_switching_templates(template_id):
    """Every optional field in SAMPLE_CONTENT has data. For every template_id
    whose spec includes that section, the data must survive into the PDF —
    the literal requirement this whole foundation exists to satisfy."""
    sections = TEMPLATE_SPECS[template_id]["sections"]
    pdf_bytes = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test Resume", sections)
    text = "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages).lower()
    if "projects" in sections:
        assert "internal analytics dashboard" in text, f"[{template_id}] lost project data"
    if "certifications" in sections:
        assert "pmp" in text, f"[{template_id}] lost certification data"
    if "achievements" in sections:
        assert "grew mau by 40" in text, f"[{template_id}] lost achievement data"
    if "languages" in sections:
        assert "hindi" in text, f"[{template_id}] lost language data"
    if "interests" in sections:
        assert "chess" in text, f"[{template_id}] lost interest data"


# ── 5. template_id resolution / validation (ephemeral-content path) ────────
# The resume_id-present path (DB is source of truth) needs a live DB row and
# is exercised by the resumes.py integration tests' download flow, not here.

@pytest.mark.asyncio
async def test_resolve_rejects_unknown_template_id_for_ephemeral_content():
    req = ExportRequest(content={}, title="Resume", resume_id=None, template_id="not-a-real-template")
    with pytest.raises(HTTPException) as exc_info:
        await _resolve_template_id(req, db=None, user=None)
    assert exc_info.value.status_code == 400


@pytest.mark.asyncio
async def test_resolve_accepts_valid_template_id_for_ephemeral_content():
    req = ExportRequest(content={}, title="Resume", resume_id=None, template_id="tech-stack")
    result = await _resolve_template_id(req, db=None, user=None)
    assert result == "tech-stack"


@pytest.mark.asyncio
async def test_resolve_defaults_to_modern_when_ephemeral_id_omitted():
    req = ExportRequest(content={}, title="Resume", resume_id=None)
    result = await _resolve_template_id(req, db=None, user=None)
    assert result == DEFAULT_TEMPLATE_ID
