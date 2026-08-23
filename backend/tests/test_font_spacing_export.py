"""
Regression tests for "Font/Spacing selected in the Resume Editor aren't
reflected in downloaded PDF/DOCX".

Source of truth: Resume.font_metadata ({"family": "sans"|"serif"|"mono",
"size": "small"|"regular"|"large"}) / Resume.layout_metadata
({"spacing": "compact"|"comfortable"|"spacious"}) — the SAME two DB columns
the Resume Editor's Font/Spacing picker already writes via PUT
/api/resumes/{id} (routers/resumes.py's ResumeUpsert), already round-tripped
through GET /api/resumes/{id} (_to_dict). This file covers the part that was
missing: the export pipeline (routers/export.py) never read them at all —
_resolve_style is the new single place that turns them into a `style` dict
every PDF/DOCX builder consumes (_apply_pdf_scale for PDF; _docx_pt_scaler/
_docx_space_scaler for DOCX).

Does NOT test pixel-level visual output (see test_template_registry.py's
docstring) — verifies the SETTING actually reaches and measurably changes
the exporter's output (font name changes, byte size changes with spacing),
and that omitting settings entirely reproduces today's exact output.
"""
import io

import pytest
from docx import Document
from pypdf import PdfReader

from routers.export import (
    TEMPLATE_BUILDERS, TEMPLATE_SPECS, _style_from_metadata,
    _FONT_SIZE_SCALE, _SPACING_SCALE, _DOCX_FONT_MAP,
)
from tests.test_template_registry import SAMPLE_CONTENT, _docx_all_text

# One classic (table/column-based) + one single-column template, so both
# families are covered without re-running every combination against all 10.
REPRESENTATIVE_IDS = ["modern", "professional", "tech-stack", "academic"]


def _pdf_text(pdf_bytes: bytes) -> str:
    return "\n".join((p.extract_text() or "") for p in PdfReader(io.BytesIO(pdf_bytes)).pages)


# ── _style_from_metadata: the merge logic itself ────────────────────────────

def test_style_from_metadata_defaults_to_template_font_when_absent():
    """No font_metadata/layout_metadata at all (a resume that's never
    touched the picker, or predates the feature) must resolve to the
    template's OWN design font and scale 1.0/1.0 — never invents a
    "sans/regular/comfortable" override that isn't actually there."""
    style = _style_from_metadata(None, None, "professional")
    assert style["font_family"] == TEMPLATE_SPECS["professional"]["font"] == "serif"
    assert style["font_scale"] == 1.0
    assert style["spacing_scale"] == 1.0


def test_style_from_metadata_honors_explicit_family_override():
    style = _style_from_metadata({"family": "mono"}, None, "professional")
    assert style["font_family"] == "mono"
    assert style["docx_font"] == _DOCX_FONT_MAP["mono"] == "Consolas"


@pytest.mark.parametrize("size_key,expected", list(_FONT_SIZE_SCALE.items()))
def test_style_from_metadata_maps_every_size_key(size_key, expected):
    style = _style_from_metadata({"size": size_key}, None, "modern")
    assert style["font_scale"] == expected


@pytest.mark.parametrize("spacing_key,expected", list(_SPACING_SCALE.items()))
def test_style_from_metadata_maps_every_spacing_key(spacing_key, expected):
    style = _style_from_metadata(None, {"spacing": spacing_key}, "modern")
    assert style["spacing_scale"] == expected


def test_style_from_metadata_ignores_garbage_family():
    """An unrecognized family value must fall back to the template's own
    design, not silently break or default to some arbitrary hardcoded font."""
    style = _style_from_metadata({"family": "comic-sans-please"}, None, "minimal")
    assert style["font_family"] == TEMPLATE_SPECS["minimal"]["font"]


# ── Font configuration reaches the PDF exporter ─────────────────────────────

@pytest.mark.parametrize("template_id", REPRESENTATIVE_IDS)
def test_font_family_reaches_pdf_exporter(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    default_bytes = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections, None)
    mono_style = _style_from_metadata({"family": "mono"}, None, template_id)
    mono_bytes = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections, mono_style)
    assert default_bytes != mono_bytes, f"[{template_id}] font override had no effect on PDF output"
    # Content survives regardless of font family.
    assert "PRIYA SHARMA" in _pdf_text(mono_bytes).upper()


# ── Font configuration reaches the DOCX exporter ────────────────────────────

@pytest.mark.parametrize("template_id", REPRESENTATIVE_IDS)
def test_font_family_reaches_docx_exporter(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    mono_style = _style_from_metadata({"family": "mono"}, None, template_id)
    docx_bytes = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test", sections, mono_style)
    doc = Document(io.BytesIO(docx_bytes))
    # Every run in the document should carry the resolved docx font name.
    font_names = {r.font.name for p in doc.paragraphs for r in p.runs if r.font.name}
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                font_names |= {r.font.name for p in cell.paragraphs for r in p.runs if r.font.name}
    assert font_names, f"[{template_id}] no runs found to check font on"
    assert font_names == {"Consolas"}, f"[{template_id}] expected only Consolas runs, found {font_names}"


# ── Spacing configuration reaches the PDF exporter ──────────────────────────

@pytest.mark.parametrize("template_id", REPRESENTATIVE_IDS)
def test_spacing_reaches_pdf_exporter(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    compact_style = _style_from_metadata(None, {"spacing": "compact"}, template_id)
    spacious_style = _style_from_metadata(None, {"spacing": "spacious"}, template_id)
    compact_bytes = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections, compact_style)
    spacious_bytes = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections, spacious_style)
    assert compact_bytes != spacious_bytes, f"[{template_id}] spacing override had no effect on PDF output"
    # Both still parse and still contain the content -- scaling never corrupts the PDF.
    for b in (compact_bytes, spacious_bytes):
        assert "PRIYA SHARMA" in _pdf_text(b).upper()


# ── Spacing configuration reaches the DOCX exporter ─────────────────────────

@pytest.mark.parametrize("template_id", REPRESENTATIVE_IDS)
def test_spacing_reaches_docx_exporter(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    compact_style = _style_from_metadata(None, {"spacing": "compact"}, template_id)
    spacious_style = _style_from_metadata(None, {"spacing": "spacious"}, template_id)
    compact_bytes = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test", sections, compact_style)
    spacious_bytes = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test", sections, spacious_style)
    assert compact_bytes != spacious_bytes, f"[{template_id}] spacing override had no effect on DOCX output"
    compact_doc = Document(io.BytesIO(compact_bytes))
    spacious_doc = Document(io.BytesIO(spacious_bytes))
    assert compact_doc.styles["Normal"].paragraph_format.line_spacing == pytest.approx(1.15 / 1.5)
    assert spacious_doc.styles["Normal"].paragraph_format.line_spacing == pytest.approx(1.9 / 1.5)
    for d in (compact_doc, spacious_doc):
        assert "PRIYA SHARMA" in _docx_all_text(d).upper()


# ── Default behavior unchanged when settings are absent ────────────────────
#
# PDF/DOCX writers embed generation-time-varying metadata in EVERY file
# they produce, entirely independent of our own code: fpdf2 writes a
# second-precision /CreationDate AND a randomized-per-generation /ID trailer
# entry; python-docx writes created/modified timestamps. Two calls with
# byte-for-byte identical content/style can therefore legitimately differ in
# their raw bytes. "Unchanged" is checked the only way that's actually
# meaningful here: identical extracted TEXT, identical measurable style
# properties (line_spacing), and byte COUNT within a few dozen bytes (a real
# font/spacing change shifts length by hundreds+ of bytes; metadata jitter
# never does).

def _pdf_page_count_and_text(pdf_bytes: bytes):
    reader = PdfReader(io.BytesIO(pdf_bytes))
    return len(reader.pages), _pdf_text(pdf_bytes)


@pytest.mark.parametrize("template_id", TEMPLATE_SPECS.keys())
def test_default_pdf_output_unchanged_when_style_omitted(template_id):
    """The literal requirement: a caller that doesn't pass `style` at all
    (every existing call site before this feature, and every resume that
    never touched the Font/Spacing picker) must get the SAME output as
    explicitly passing None/{} — see module note on why "same" means
    content-identical rather than byte-identical for PDF/DOCX."""
    sections = TEMPLATE_SPECS[template_id]["sections"]
    no_arg = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections)
    explicit_none = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections, None)
    empty_dict = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections, {})
    results = [_pdf_page_count_and_text(b) for b in (no_arg, explicit_none, empty_dict)]
    assert results[0] == results[1] == results[2]
    lengths = [len(no_arg), len(explicit_none), len(empty_dict)]
    assert max(lengths) - min(lengths) < 30, f"[{template_id}] byte-length drift suggests a real content/style change, not just metadata jitter: {lengths}"


@pytest.mark.parametrize("template_id", TEMPLATE_SPECS.keys())
def test_default_docx_output_unchanged_when_style_omitted(template_id):
    sections = TEMPLATE_SPECS[template_id]["sections"]
    no_arg = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test", sections)
    explicit_none = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test", sections, None)
    empty_dict = TEMPLATE_BUILDERS[template_id].docx(SAMPLE_CONTENT, "Test", sections, {})
    docs = [Document(io.BytesIO(b)) for b in (no_arg, explicit_none, empty_dict)]
    texts = [_docx_all_text(d) for d in docs]
    assert texts[0] == texts[1] == texts[2]
    line_spacings = [d.styles["Normal"].paragraph_format.line_spacing for d in docs]
    assert line_spacings[0] == line_spacings[1] == line_spacings[2]
    lengths = [len(no_arg), len(explicit_none), len(empty_dict)]
    assert max(lengths) - min(lengths) < 30, f"[{template_id}] byte-length drift suggests a real content/style change, not just metadata jitter: {lengths}"


def test_default_style_from_metadata_matches_omitted_style_exactly():
    """_style_from_metadata(None, None, tid) — what a resume with no saved
    font/layout metadata resolves to — must produce a style dict that is a
    complete no-op (scale 1.0/1.0), so routing it through the builders
    changes nothing vs. omitting style entirely."""
    for template_id in TEMPLATE_SPECS:
        sections = TEMPLATE_SPECS[template_id]["sections"]
        style = _style_from_metadata(None, None, template_id)
        assert style["font_scale"] == 1.0 and style["spacing_scale"] == 1.0
        via_style = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections, style)
        omitted = TEMPLATE_BUILDERS[template_id].pdf(SAMPLE_CONTENT, "Test", sections)
        assert _pdf_page_count_and_text(via_style) == _pdf_page_count_and_text(omitted), \
            f"[{template_id}] resolved default style diverged from omitted style"
        assert abs(len(via_style) - len(omitted)) < 30


# ── Project + Custom Section + template still survive under font/spacing ───

@pytest.mark.parametrize("template_id", REPRESENTATIVE_IDS)
def test_project_and_custom_section_and_template_survive_font_and_spacing(template_id):
    content = dict(SAMPLE_CONTENT)
    content["customSections"] = [{"id": "cs1", "title": "Professional Highlights",
                                   "content": "Client support, UAT, functional testing."}]
    sections = TEMPLATE_SPECS[template_id]["sections"]
    style = _style_from_metadata({"family": "mono", "size": "large"}, {"spacing": "compact"}, template_id)
    pdf_text = _pdf_text(TEMPLATE_BUILDERS[template_id].pdf(content, "Test", sections, style)).lower()
    docx_text = _docx_all_text(Document(io.BytesIO(
        TEMPLATE_BUILDERS[template_id].docx(content, "Test", sections, style)))).lower()
    assert "internal analytics dashboard" in pdf_text, f"[{template_id}] PDF lost project under font/spacing"
    assert "professional highlights" in pdf_text, f"[{template_id}] PDF lost custom section under font/spacing"
    assert "internal analytics dashboard" in docx_text, f"[{template_id}] DOCX lost project under font/spacing"
    assert "professional highlights" in docx_text, f"[{template_id}] DOCX lost custom section under font/spacing"
