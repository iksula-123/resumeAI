"""
Resume export — PDF and DOCX, template-aware.

template_id resolution (ONE source of truth, never mixed):
  - resume_id given  → the resume's OWN Resume.template_id (DB) wins. Any
                        client-supplied template_id in the request body is
                        ignored outright — a saved resume's download must
                        always match what's actually saved, never a client
                        guess that could've gone stale.
  - resume_id absent → this is ephemeral/unsaved content (e.g. the template
                        picker's "preview with my data" or AI-Upgrade before
                        the user has saved anything yet). The client-supplied
                        template_id is the only signal available, so it's
                        used — but strictly validated against
                        shared/template-specs.json; an unknown id is a 400,
                        not a silent fallback (see TEMPLATE-VALIDATION below).

This is completely separate from the uploaded-resume "preserve original
design" workflow (Resume.preserve_original / template_type == "uploaded_original"
— see routers/resumes.py::download_resume and services/docx_editor.py).
template_id only matters when preserve_original is false; the two concepts
are never read together.

TEMPLATE_BUILDERS is a registry, not 10 independent implementations — every
template_id currently maps to the same "classic" layout-family builder
(today's fpdf2/python-docx design, now parameterized by which sections to
render). As real per-template visual designs get built, new layout-family
builder functions get added and specific ids get repointed to them — the
registry shape doesn't change.
"""
import io
import json
import os
import re
import uuid
from dataclasses import dataclass
from typing import Callable

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from database import get_db
from models import Resume
from services.auth import verify_token
from services.storage import upload_bytes
from services.webhooks import dispatch
from services.usage import log_usage_event

router = APIRouter(prefix="/api/export", tags=["Export"])
security = HTTPBearer()


async def _auth(c: HTTPAuthorizationCredentials = Depends(security)):
    return await verify_token(c.credentials)


class ExportRequest(BaseModel):
    content: dict
    title: str = "Resume"
    resume_id: str | None = None
    template_id: str = "modern"  # only consulted when resume_id is absent — see module docstring
    # font_metadata: {"family": "sans"|"serif"|"mono", "size": "small"|"regular"|"large"}
    # layout_metadata: {"spacing": "compact"|"comfortable"|"spacious"}
    # Same shape, same source of truth as Resume.font_metadata/layout_metadata
    # (the Resume Editor's Font/Spacing picker — see routers/resumes.py's
    # ResumeUpsert). Only consulted when resume_id is absent — see
    # _resolve_style's docstring for why (mirrors template_id's own rule).
    font_metadata: dict | None = None
    layout_metadata: dict | None = None


# ── shared/template-specs.json: the single source of truth ────────────────────
# Same file frontend/components/ResumeTemplates.tsx reads. Loaded once at
# import time; do not hand-maintain a second copy of this data here.

def _load_template_specs() -> dict[str, dict]:
    path = os.path.join(os.path.dirname(__file__), "..", "..", "shared", "template-specs.json")
    with open(path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    return {t["id"]: t for t in raw["templates"]}


TEMPLATE_SPECS: dict[str, dict] = _load_template_specs()
DEFAULT_TEMPLATE_ID = "modern"


def _skill_names(content: dict) -> list[str]:
    """Skills may be objects {name, level} or plain strings."""
    out = []
    for s in content.get("skills", []) or []:
        if isinstance(s, dict):
            name = s.get("name")
        else:
            name = s
        if name:
            out.append(str(name))
    return out


# fpdf2 core fonts only support latin-1; map common unicode to safe equivalents
_UNICODE_MAP = {
    "–": "-", "—": "-", "•": "-",   # en/em dash, bullet
    "‘": "'", "’": "'", "“": '"', "”": '"',  # smart quotes
    "…": "...", " ": " ",
}


def _pdf_safe(text) -> str:
    text = str(text or "")
    for uni, rep in _UNICODE_MAP.items():
        text = text.replace(uni, rep)
    # drop anything still outside latin-1 so fpdf never raises
    return text.encode("latin-1", "replace").decode("latin-1")


_FILENAME_UNSAFE_RE = re.compile(r'[\\/:*?"<>|\s]+')


def _safe_filename_base(title: str) -> str:
    """A resume title used as a downloaded filename needs more than just
    spaces replaced — titles like the auto-generated "Full Name — Job Title
    / Other Title" (the ATS Checker's "Improve My Resume" flow's own
    pattern) contain '/' and other characters that are invalid or act as
    path separators on common filesystems."""
    base = _FILENAME_UNSAFE_RE.sub("_", title or "Resume").strip("_")
    return base or "Resume"


def _content_disposition(filename: str) -> str:
    """Builds a Content-Disposition header value that's safe for ANY resume
    title. HTTP header values must be latin-1, but resume titles routinely
    contain characters outside it — e.g. the "Full Name — Job Title" title
    the ATS Checker's "Improve My Resume" flow auto-generates contains an
    em dash (U+2014). Building the header from that title directly used to
    raise UnicodeEncodeError inside Starlette's response init and turn EVERY
    download for that resume into a 500, regardless of which template was
    selected. Encodes both a latin-1-safe fallback (`filename=`, legacy
    clients) and the RFC 5987 UTF-8 form (`filename*=`, what every modern
    browser actually displays/saves as) so non-ASCII titles download
    correctly instead of crashing the response."""
    from urllib.parse import quote
    ascii_fallback = filename.encode("ascii", "ignore").decode("ascii").strip() or "resume"
    utf8_encoded = quote(filename, safe="")
    return f'attachment; filename="{ascii_fallback}"; filename*=UTF-8\'\'{utf8_encoded}'


def _hex_to_rgb(hex_color: str | None) -> tuple[int, int, int]:
    h = (hex_color or "").lstrip("#")
    if len(h) != 6:
        return (30, 30, 30)
    try:
        return (int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    except ValueError:
        return (30, 30, 30)


def _mix_rgb(a: tuple[int, int, int], b: tuple[int, int, int], t: float) -> tuple[int, int, int]:
    """Linear-interpolate between two RGB colors (t=0 -> a, t=1 -> b). Used to
    approximate a CSS gradient with a band of solid-color strips, since fpdf2
    has no native gradient fill."""
    return tuple(round(a[i] + (b[i] - a[i]) * t) for i in range(3))


# ── template_id resolution (the one source-of-truth chokepoint) ────────────────

async def _resolve_template_id(req: ExportRequest, db: AsyncSession, user) -> str:
    if req.resume_id:
        try:
            rid = uuid.UUID(req.resume_id)
        except ValueError:
            raise HTTPException(status_code=404, detail="Resume not found")
        row = (await db.execute(
            select(Resume.template_id, Resume.user_id).where(Resume.id == rid)
        )).first()
        # NOTE: this router's _auth/verify_token returns the raw Supabase auth
        # user (id is a str), unlike services/deps.py::get_current_user (used
        # elsewhere) which returns our Profile row (id is a uuid.UUID) — so
        # this comparison must be string-normalized on both sides, or a UUID
        # vs str mismatch makes every resume look "not owned" by its own owner.
        if not row or (str(row.user_id) != str(user.id) and not getattr(user, "is_admin", False)):
            raise HTTPException(status_code=404, detail="Resume not found")
        tid = row.template_id or DEFAULT_TEMPLATE_ID
        if tid not in TEMPLATE_SPECS:
            # Backward compatibility ONLY: a resume already saved with a
            # legacy/removed template_id must still download, never 400 —
            # the user did nothing wrong, the id just predates this registry.
            tid = DEFAULT_TEMPLATE_ID
        return tid

    # No resume_id — ephemeral content (template picker preview, AI-Upgrade
    # pre-save export). The client-supplied id is the only signal, and unlike
    # the DB-sourced path above, an unknown id here is NOT quietly patched to
    # "modern" — that would let a typo'd/invalid id silently render the wrong
    # template with no signal to the caller.
    tid = req.template_id or DEFAULT_TEMPLATE_ID
    if tid not in TEMPLATE_SPECS:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown template_id '{tid}'. Valid ids: {sorted(TEMPLATE_SPECS)}",
        )
    return tid


# Same three values the Font picker's Size row offers / the Spacing picker
# offers (frontend/app/resumes/[id]/edit/page.tsx) — kept as the ONE mapping
# from those keys to a numeric scale, mirrored from the exact multipliers
# the editor's own Preview pane uses (`zoom:` / `line-height:` in its scoped
# <style> override) so PDF/DOCX spacing/size actually matches Preview
# instead of inventing separate numbers. Spacing is normalized so
# "comfortable" (this exporter's own unscaled baseline, tuned before this
# feature existed) stays neutral at 1.0, preserving the SAME relative
# tightening/loosening ratio the frontend uses relative to its own
# comfortable baseline (1.5).
_FONT_SIZE_SCALE = {"small": 0.85, "regular": 1.0, "large": 1.15}
_SPACING_SCALE = {"compact": 1.15 / 1.5, "comfortable": 1.0, "spacious": 1.9 / 1.5}
_DOCX_FONT_MAP = {"sans": "Calibri", "serif": "Georgia", "mono": "Consolas"}


def _style_from_metadata(font_metadata: dict | None, layout_metadata: dict | None, template_id: str) -> dict:
    """The one place that turns font_metadata/layout_metadata (whatever
    their source — see _resolve_style and routers/resumes.py::download_resume,
    the two callers) into the flat style dict every builder consumes.
    font_family always has a value (the override, or the template's OWN
    design default from shared/template-specs.json), while font_scale/
    spacing_scale default to 1.0 (today's unscaled behavior) whenever no
    explicit size/spacing was ever saved — the literal "default behavior
    remains unchanged when settings are absent" requirement."""
    spec = TEMPLATE_SPECS.get(template_id, {})
    family = (font_metadata or {}).get("family")
    if family not in ("sans", "serif", "mono"):
        family = spec.get("font") or "sans"
    size_key = (font_metadata or {}).get("size")
    spacing_key = (layout_metadata or {}).get("spacing")
    return {
        "font_family": family,
        "font_scale": _FONT_SIZE_SCALE.get(size_key, 1.0),
        "spacing_scale": _SPACING_SCALE.get(spacing_key, 1.0),
        "docx_font": _DOCX_FONT_MAP.get(family, "Calibri"),
    }


async def _resolve_style(req: ExportRequest, db: AsyncSession, user, template_id: str) -> dict:
    """Resolves font_metadata/layout_metadata the SAME way _resolve_template_id
    resolves template_id: a saved resume's OWN Resume.font_metadata/
    layout_metadata (DB) wins over anything client-supplied — a download must
    always match what's actually saved, never a stale client guess. For
    ephemeral content (no resume_id) the client-supplied values are the only
    signal. Ownership is already checked by the _resolve_template_id call
    every caller makes first, so this doesn't repeat that check."""
    font_metadata: dict | None = None
    layout_metadata: dict | None = None
    if req.resume_id:
        try:
            rid = uuid.UUID(req.resume_id)
        except ValueError:
            rid = None
        if rid is not None:
            row = (await db.execute(
                select(Resume.font_metadata, Resume.layout_metadata).where(Resume.id == rid)
            )).first()
            if row:
                font_metadata, layout_metadata = row.font_metadata, row.layout_metadata
    else:
        font_metadata, layout_metadata = req.font_metadata, req.layout_metadata

    return _style_from_metadata(font_metadata, layout_metadata, template_id)


def _apply_pdf_scale(pdf, font_scale: float, spacing_scale: float):
    """Scales every font-size and every spacing/line-height call a PDF
    builder function makes, WITHOUT touching each individual call site —
    wraps the fpdf2 primitives that carry those two concerns (set_font's
    size arg; ln/cell/multi_cell's height arg) on THIS pdf instance only.
    font_scale comes from the Font picker's Size row (small/regular/large);
    spacing_scale from the Spacing picker (compact/comfortable/spacious) —
    see _resolve_style. A no-op when both are 1.0 (today's exact behavior,
    byte-for-byte, since the wrapped functions are never installed)."""
    if font_scale == 1.0 and spacing_scale == 1.0:
        return
    _orig_set_font = pdf.set_font
    _orig_ln = pdf.ln
    _orig_multi_cell = pdf.multi_cell
    _orig_cell = pdf.cell

    def set_font(family=None, style="", size=0, **kw):
        return _orig_set_font(family, style, (size * font_scale) if size else size, **kw)

    def ln(h=None):
        return _orig_ln(h * spacing_scale) if h is not None else _orig_ln(h)

    def multi_cell(w, h=None, *a, **kw):
        return _orig_multi_cell(w, (h * spacing_scale) if h else h, *a, **kw)

    def cell(w=None, h=None, *a, **kw):
        return _orig_cell(w, (h * spacing_scale) if h else h, *a, **kw)

    pdf.set_font, pdf.ln, pdf.multi_cell, pdf.cell = set_font, ln, multi_cell, cell


def _docx_pt_scaler(font_scale: float):
    """A `Pt`-shaped constructor pre-scaled by font_scale, meant to shadow
    the real `Pt` import (`Pt = _docx_pt_scaler(font_scale)`) for the rest
    of a DOCX builder function — every subsequent `Pt(N)` FONT SIZE (and
    margin) call scales automatically, with no per-call-site edits needed.
    Scaling margins along with font size too mirrors the editor Preview's
    own `zoom:` CSS property (the Font picker's Size row), which scales the
    ENTIRE rendered box proportionally, not just glyph size."""
    from docx.shared import Pt as _RealPt
    return lambda n: _RealPt(n * font_scale)


def _docx_space_scaler(spacing_scale: float):
    """A `Pt`-shaped constructor pre-scaled by spacing_scale, for
    paragraph_format.space_before/space_after specifically (section/entry
    gaps) — kept separate from _docx_pt_scaler's font_scale because they
    come from two different pickers (Font Size vs Spacing) and must be able
    to vary independently, exactly like the two are independent in the
    editor's own Font/Spacing pickers."""
    from docx.shared import Pt as _RealPt
    return lambda n: _RealPt(n * spacing_scale)


@router.post("/pdf")
async def export_pdf(req: ExportRequest, db: AsyncSession = Depends(get_db), user=Depends(_auth)):
    template_id = await _resolve_template_id(req, db, user)
    style = await _resolve_style(req, db, user, template_id)
    builder = TEMPLATE_BUILDERS.get(template_id, TEMPLATE_BUILDERS[DEFAULT_TEMPLATE_ID])
    sections = TEMPLATE_SPECS[template_id]["sections"]

    pdf_bytes = builder.pdf(req.content, req.title, sections, style)
    safe_title = _safe_filename_base(req.title)
    upload_bytes(str(user.id), "generated", f"{safe_title}.pdf", pdf_bytes, "application/pdf")
    dispatch(user.id, "resume.exported", {"title": req.title, "format": "pdf", "template_id": template_id})
    await log_usage_event(str(user.id), "download_pdf", metadata={"title": req.title, "template_id": template_id})
    return StreamingResponse(
        io.BytesIO(pdf_bytes),
        media_type="application/pdf",
        headers={"Content-Disposition": _content_disposition(f"{safe_title}.pdf")},
    )


@router.post("/docx")
async def export_docx(req: ExportRequest, db: AsyncSession = Depends(get_db), user=Depends(_auth)):
    template_id = await _resolve_template_id(req, db, user)
    style = await _resolve_style(req, db, user, template_id)
    builder = TEMPLATE_BUILDERS.get(template_id, TEMPLATE_BUILDERS[DEFAULT_TEMPLATE_ID])
    sections = TEMPLATE_SPECS[template_id]["sections"]

    docx_bytes = builder.docx(req.content, req.title, sections, style)
    safe_title = _safe_filename_base(req.title)
    upload_bytes(
        str(user.id), "generated", f"{safe_title}.docx", docx_bytes,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    dispatch(user.id, "resume.exported", {"title": req.title, "format": "docx", "template_id": template_id})
    await log_usage_event(str(user.id), "download_docx", metadata={"title": req.title, "template_id": template_id})
    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _content_disposition(f"{safe_title}.docx")},
    )


# ──────────────────────────────────────────────
# PDF generation via fpdf2 — the "classic" layout family
# (today's Modern design, now parameterized by `sections`)
# ──────────────────────────────────────────────

_PDF_FONT_FAMILY_MAP = {"sans": "NotoSans", "serif": "NotoSerif", "mono": "NotoSansMono"}
_PDF_CORE_FALLBACK_MAP = {"sans": "Helvetica", "serif": "Times", "mono": "Courier"}


def _pdf_font(pdf, family: str = "sans"):
    """Set up Unicode fonts for the PDF and return (font_name, safe_fn).

    `family` ("sans" | "serif" | "mono") — the SAME three values the Resume
    Editor's own Font picker offers (frontend/app/resumes/[id]/edit/page.tsx)
    — lets a template's own typeface choice, or the candidate's explicit
    override (see _resolve_style/font_metadata), actually reach the PDF
    instead of every template always rendering Noto Sans regardless of
    design. Whichever primary font loads, Noto Sans Devanagari is always
    registered as a *fallback*, so English AND Devanagari/Indian scripts
    render correctly (spec Section 6) no matter which primary typeface is
    active. Falls back to a latin-1 core font (Helvetica/Times/Courier) when
    the TTFs aren't bundled. Drop the OFL TTFs into backend/assets/fonts/.
    """
    base = os.path.join(os.path.dirname(__file__), "..", "assets", "fonts")
    font_name = _PDF_FONT_FAMILY_MAP.get(family, "NotoSans")
    core_fallback = _PDF_CORE_FALLBACK_MAP.get(family, "Helvetica")

    latin_reg = os.path.join(base, f"{font_name}-Regular.ttf")
    if not os.path.exists(latin_reg):
        # Requested family's TTF isn't bundled -- try the sans TTF before
        # giving up entirely, then the matching core (latin-1 only) font.
        latin_reg = os.path.join(base, "NotoSans-Regular.ttf")
        if not os.path.exists(latin_reg):
            return core_fallback, _pdf_safe
        font_name = "NotoSans"

    latin_bold = os.path.join(base, f"{font_name}-Bold.ttf")
    pdf.add_font(font_name, "", latin_reg)
    pdf.add_font(font_name, "B", latin_bold if os.path.exists(latin_bold) else latin_reg)

    # Devanagari as a fallback so mixed English/Hindi text renders in one run,
    # regardless of whether the primary font is Noto Sans or Noto Serif.
    dev_reg = os.path.join(base, "NotoSansDevanagari-Regular.ttf")
    if os.path.exists(dev_reg):
        dev_bold = os.path.join(base, "NotoSansDevanagari-Bold.ttf")
        pdf.add_font("NotoDev", "", dev_reg)
        pdf.add_font("NotoDev", "B", dev_bold if os.path.exists(dev_bold) else dev_reg)
        try:
            pdf.set_fallback_fonts(["NotoDev"])
        except Exception:
            pass
    return font_name, (lambda t: str(t or ""))   # pass Unicode straight through


# Historical default — exactly what _build_pdf/_build_docx rendered before
# they became sections-aware. Only used as a safety net for a caller that
# doesn't go through TEMPLATE_BUILDERS/TEMPLATE_SPECS (e.g. a direct unit
# test); every real request path always passes an explicit `sections` list
# resolved from the selected template's spec.
_LEGACY_DEFAULT_SECTIONS = ["summary", "experience", "education", "skills"]


def _build_pdf(content: dict, title: str, sections: list[str] | None = None) -> bytes:
    """The 'classic' PDF layout family. `sections` (from the resolved
    template's spec — see TEMPLATE_SPECS) is the ONLY thing that decides which
    optional blocks render; this function never makes that call independently
    (spec Section 7 — section-parity single source of truth). A section with
    no data is always skipped regardless of `sections`, so nothing empty ever
    renders — and nothing with data is dropped just because a different
    template got selected."""
    sections = sections if sections is not None else _LEGACY_DEFAULT_SECTIONS
    from fpdf import FPDF

    pi = content.get("personalInfo", {})

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    font, safe = _pdf_font(pdf)   # Unicode-capable when the font is bundled

    # Name
    pdf.set_font(font, "B", 20)
    pdf.set_text_color(30, 64, 175)
    pdf.cell(0, 10, safe(pi.get("fullName", title)), ln=True)

    # Contact line
    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location")] if s]
    if contact_parts:
        pdf.set_font(font, "", 9)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 6, safe("  |  ".join(contact_parts)), ln=True)

    pdf.ln(3)

    def section(heading: str):
        pdf.set_font(font, "B", 10)
        pdf.set_text_color(30, 64, 175)
        pdf.set_fill_color(239, 246, 255)
        pdf.cell(0, 7, safe(heading.upper()), ln=True, fill=True)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font(font, "", 10)

    def body_line(text: str, size: int = 10, bold: bool = False, grey: bool = False):
        pdf.set_x(pdf.l_margin)  # always start at left margin so width is never 0
        pdf.set_font(font, "B" if bold else "", size)
        pdf.set_text_color(120, 120, 120) if grey else pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(0, 5 if not bold else 6, safe(text))
        pdf.set_text_color(30, 30, 30)

    if "summary" in sections and content.get("summary"):
        section("Summary")
        body_line(content["summary"])
        pdf.ln(4)

    if "experience" in sections and content.get("experience"):
        section("Experience")
        for exp in content["experience"]:
            position = exp.get("position", "")
            company = exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            date_str = f"  ({start} - {end})" if start else ""
            label = position + (f" - {company}" if company else "") + date_str
            body_line(label, bold=True)
            for bullet in exp.get("bullets", []):
                if str(bullet).strip():
                    body_line(f"  -  {bullet}")
            pdf.ln(2)

    if "education" in sections and content.get("education"):
        section("Education")
        for edu in content["education"]:
            inst = edu.get("institution", "")
            deg = ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            end_date = edu.get("endDate", "")
            body_line(inst + (f"  ({end_date})" if end_date else ""), bold=True)
            if deg:
                body_line(deg)
            pdf.ln(2)

    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            section("Skills")
            body_line("  -  ".join(skill_names))
            pdf.ln(2)

    if "projects" in sections and content.get("projects"):
        section("Projects")
        for proj in content["projects"]:
            name = proj.get("name", "")
            tech = proj.get("technologies", "")
            label = name + (f"  ({tech})" if tech else "")
            if label.strip():
                body_line(label, bold=True)
            if proj.get("description"):
                body_line(proj["description"])
            pdf.ln(1)
        pdf.ln(1)

    if "certifications" in sections and content.get("certifications"):
        section("Certifications")
        for cert in content["certifications"]:
            name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            line = name + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                body_line(line)
        pdf.ln(2)

    if "achievements" in sections and content.get("achievements"):
        section("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                body_line(f"  -  {a}")
        pdf.ln(2)

    if "languages" in sections and content.get("languages"):
        lang_strs = []
        for l in content["languages"]:
            if isinstance(l, dict):
                nm, prof = l.get("name", ""), l.get("proficiency", "")
                if nm:
                    lang_strs.append(f"{nm} ({prof})" if prof else nm)
            elif str(l).strip():
                lang_strs.append(str(l))
        if lang_strs:
            section("Languages")
            body_line("  -  ".join(lang_strs))
            pdf.ln(2)

    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            section("Interests")
            body_line("  -  ".join(interests))
        pdf.ln(2)

    # Custom (user-named) sections -- always rendered last, matching the
    # Resume Builder's own "Resume Sections" panel order (see
    # frontend/app/resumes/[id]/edit/page.tsx and CustomSectionsBlock.tsx).
    # The section TITLE is always whatever the candidate actually typed --
    # never hardcoded, never defaulted to a specific example name.
    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title = (cs.get("title") or "").strip()
            cs_body = (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            section(cs_title or "Additional Information")
            if cs_body:
                body_line(cs_body)
            pdf.ln(1)

    return bytes(pdf.output())


# ──────────────────────────────────────────────
# DOCX generation via python-docx — the "classic" layout family
# ──────────────────────────────────────────────

def _build_docx(content: dict, title: str, sections: list[str] | None = None) -> bytes:
    """DOCX counterpart of _build_pdf — same `sections`-driven contract."""
    sections = sections if sections is not None else _LEGACY_DEFAULT_SECTIONS
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pi = content.get("personalInfo", {})
    doc = Document()

    # Remove default margins slightly
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(36)
        sec.left_margin = sec.right_margin = Pt(54)

    # Name heading
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = name_para.add_run(pi.get("fullName", title))
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    # Contact
    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location")] if s]
    if contact_parts:
        cp = doc.add_paragraph("  |  ".join(contact_parts))
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    def add_section(heading: str):
        h = doc.add_paragraph(heading.upper())
        run = h.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        h.paragraph_format.space_before = Pt(10)

    if "summary" in sections and content.get("summary"):
        add_section("Summary")
        doc.add_paragraph(content["summary"])

    if "experience" in sections and content.get("experience"):
        add_section("Experience")
        for exp in content["experience"]:
            p = doc.add_paragraph()
            r = p.add_run(f"{exp.get('position', '')} — {exp.get('company', '')}")
            r.font.bold = True
            for bullet in exp.get("bullets", []):
                if bullet.strip():
                    bp = doc.add_paragraph(bullet, style="List Bullet")
                    bp.paragraph_format.left_indent = Pt(18)

    if "education" in sections and content.get("education"):
        add_section("Education")
        for edu in content["education"]:
            p = doc.add_paragraph()
            r = p.add_run(edu.get("institution", ""))
            r.font.bold = True
            deg = ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            if deg:
                doc.add_paragraph(deg)

    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            add_section("Skills")
            doc.add_paragraph("  •  ".join(skill_names))

    if "projects" in sections and content.get("projects"):
        add_section("Projects")
        for proj in content["projects"]:
            name = proj.get("name", "")
            tech = proj.get("technologies", "")
            if name or tech:
                p = doc.add_paragraph()
                r = p.add_run(name + (f"  ({tech})" if tech else ""))
                r.font.bold = True
            if proj.get("description"):
                doc.add_paragraph(proj["description"])

    if "certifications" in sections and content.get("certifications"):
        add_section("Certifications")
        for cert in content["certifications"]:
            name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            line = name + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                doc.add_paragraph(line)

    if "achievements" in sections and content.get("achievements"):
        add_section("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                doc.add_paragraph(str(a), style="List Bullet")

    if "languages" in sections and content.get("languages"):
        lang_strs = []
        for l in content["languages"]:
            if isinstance(l, dict):
                nm, prof = l.get("name", ""), l.get("proficiency", "")
                if nm:
                    lang_strs.append(f"{nm} ({prof})" if prof else nm)
            elif str(l).strip():
                lang_strs.append(str(l))
        if lang_strs:
            add_section("Languages")
            doc.add_paragraph("  •  ".join(lang_strs))

    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            add_section("Interests")
            doc.add_paragraph("  •  ".join(interests))

    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title = (cs.get("title") or "").strip()
            cs_body = (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            add_section(cs_title or "Additional Information")
            if cs_body:
                doc.add_paragraph(cs_body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────
# "single-column" layout family — Phase 2's five new templates
# (tech-stack, fresher, academic, healthcare, international)
#
# One parameterized PDF function + one parameterized DOCX function shared by
# all five — genuinely distinct output per template via SINGLE_COLUMN_CONFIGS
# (accent color, section heading labels, section order, skill grouping,
# certification emphasis), not five independent implementations. Section
# order/labels here MUST mirror the matching React component exactly (see
# frontend/components/templates/*.tsx docstrings) — that agreement IS the
# section-parity contract; there's no runtime link between them, so keep
# them in sync by hand when either changes.
# ──────────────────────────────────────────────

def _fresher_section_order(has_experience: bool) -> list[str]:
    """Adaptive order (mirrors FresherTemplate.tsx): no large empty
    Experience block — Education and Projects promote when there's none.
    "customSections" is always last, matching the Resume Builder's own
    section panel order — never inserted mid-resume."""
    if has_experience:
        return ["summary", "skills", "experience", "education", "projects",
                 "certifications", "achievements", "interests", "languages", "customSections"]
    return ["summary", "education", "projects", "skills",
             "certifications", "achievements", "interests", "languages", "customSections"]


SINGLE_COLUMN_CONFIGS: dict[str, dict] = {
    "tech-stack": {
        "accent": (14, 165, 233), "docx_font": "Calibri",
        "labels": {"summary": "Professional Summary", "skills": "Technical Skills", "experience": "Experience",
                   "projects": "Projects", "certifications": "Certifications", "education": "Education",
                   "achievements": "Achievements", "languages": "Languages", "interests": "Interests"},
        "order": ["summary", "skills", "experience", "projects", "certifications",
                   "education", "achievements", "languages", "interests", "customSections"],
        "group_skills": True,
    },
    "fresher": {
        "accent": (22, 163, 74), "docx_font": "Calibri",
        "labels": {"summary": "Career Objective", "skills": "Technical Skills", "experience": "Internships & Experience",
                   "education": "Education", "projects": "Projects", "certifications": "Certifications",
                   "achievements": "Achievements", "interests": "Extracurricular Activities", "languages": "Languages"},
        "order": _fresher_section_order,
        "group_skills": False,
    },
    "academic": {
        "accent": (124, 45, 18), "docx_font": "Georgia",
        "labels": {"summary": "Academic Profile", "education": "Education", "skills": "Areas of Expertise",
                   "experience": "Academic & Research Experience", "projects": "Publications & Research Projects",
                   "achievements": "Awards & Honors", "certifications": "Certifications & Professional Development",
                   "languages": "Languages", "interests": "Interests"},
        "order": ["summary", "education", "skills", "experience", "projects",
                   "achievements", "certifications", "languages", "interests", "customSections"],
        "group_skills": False,
    },
    "healthcare": {
        "accent": (13, 148, 136), "docx_font": "Calibri",
        "labels": {"summary": "Professional Profile", "experience": "Clinical Experience",
                   "certifications": "Licenses & Certifications", "skills": "Clinical Skills", "education": "Education",
                   "projects": "Clinical Training & Projects", "achievements": "Achievements",
                   "languages": "Languages", "interests": "Interests"},
        "order": ["summary", "experience", "certifications", "skills", "education",
                   "projects", "achievements", "languages", "interests", "customSections"],
        "group_skills": False, "emphasize_certifications": True,
    },
    "international": {
        "accent": (30, 41, 59), "docx_font": "Georgia",
        "labels": {"summary": "Professional Summary", "skills": "Core Competencies", "experience": "Professional Experience",
                   "achievements": "Key Achievements", "education": "Education", "certifications": "Certifications",
                   "projects": "Projects", "languages": "Languages", "interests": "Additional Information"},
        "order": ["summary", "skills", "experience", "achievements", "education",
                   "certifications", "projects", "languages", "interests", "customSections"],
        "group_skills": False,
    },
}


def _language_strings(content: dict) -> list[str]:
    out = []
    for l in content.get("languages") or []:
        if isinstance(l, dict):
            nm, prof = l.get("name", ""), l.get("proficiency", "")
            if nm:
                out.append(f"{nm} ({prof})" if prof else nm)
        elif str(l).strip():
            out.append(str(l))
    return out


def _build_pdf_single_column(content: dict, title: str, sections: list[str], *, template_id: str,
                              style: dict | None = None) -> bytes:
    """`style` (see _resolve_style) carries the candidate's own explicit
    font/spacing override on top of this template's own design; absent/None
    means "use the template's own design, no scaling" — today's exact
    behavior, byte-for-byte."""
    from fpdf import FPDF
    from services.skill_categories import group_skills_by_category

    config = SINGLE_COLUMN_CONFIGS[template_id]
    accent = config["accent"]
    labels = config["labels"]
    order_spec = config["order"]
    order = order_spec(bool(content.get("experience"))) if callable(order_spec) else order_spec

    style = style or {}
    font_family = style.get("font_family") or TEMPLATE_SPECS.get(template_id, {}).get("font") or "sans"
    font_scale = style.get("font_scale", 1.0)
    spacing_scale = style.get("spacing_scale", 1.0)

    pi = content.get("personalInfo", {})
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    font, safe = _pdf_font(pdf, font_family)
    _apply_pdf_scale(pdf, font_scale, spacing_scale)

    pdf.set_font(font, "B", 20)
    pdf.set_text_color(*accent)
    pdf.cell(0, 10, safe(pi.get("fullName", title)), ln=True)

    if pi.get("jobTitle"):
        pdf.set_font(font, "", 10)
        pdf.set_text_color(90, 90, 90)
        pdf.cell(0, 6, safe(pi["jobTitle"]), ln=True)

    contact_parts = [s for s in [pi.get("location"), pi.get("phone"), pi.get("email"),
                                    pi.get("linkedin"), pi.get("github"), pi.get("website")] if s]
    if contact_parts:
        pdf.set_font(font, "", 9)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 6, safe("  |  ".join(contact_parts)), ln=True)
    pdf.ln(3)
    pdf.set_text_color(30, 30, 30)

    def section(heading: str):
        pdf.set_font(font, "B", 10)
        pdf.set_text_color(*accent)
        pdf.cell(0, 7, safe(heading.upper()), ln=True)
        pdf.set_draw_color(*accent)
        pdf.set_line_width(0.4)
        y = pdf.get_y()
        pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
        pdf.ln(2)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font(font, "", 10)

    def body_line(text: str, size: int = 10, bold: bool = False, fill: bool = False):
        pdf.set_x(pdf.l_margin)
        pdf.set_font(font, "B" if bold else "", size)
        pdf.set_text_color(30, 30, 30)
        if fill:
            light = tuple(min(255, c + 190) for c in accent)
            pdf.set_fill_color(*light)
        pdf.multi_cell(0, 5 if not bold else 6, safe(text), fill=fill)

    def render_summary():
        if content.get("summary"):
            section(labels["summary"]); body_line(content["summary"]); pdf.ln(3)

    def render_skills():
        names = _skill_names(content)
        if not names:
            return
        section(labels["skills"])
        if config.get("group_skills"):
            for cat, cat_skills in group_skills_by_category(names):
                body_line(f"{cat}: {', '.join(cat_skills)}")
        else:
            body_line("  ·  ".join(names))
        pdf.ln(3)

    def render_experience():
        if not content.get("experience"):
            return
        section(labels["experience"])
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            date_str = f"  ({start} - {end})" if start else ""
            body_line(position + (f" - {company}" if company else "") + date_str, bold=True)
            for bullet in exp.get("bullets", []):
                if str(bullet).strip():
                    body_line(f"  -  {bullet}")
            pdf.ln(2)

    def render_education():
        if not content.get("education"):
            return
        section(labels["education"])
        for edu in content["education"]:
            inst, deg = edu.get("institution", ""), ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            end_date = edu.get("endDate", "")
            body_line((deg or inst) + (f"  ({end_date})" if end_date else ""), bold=True)
            if deg and inst:
                body_line(inst)
            pdf.ln(2)

    def render_projects():
        if not content.get("projects"):
            return
        section(labels["projects"])
        for proj in content["projects"]:
            name, tech = proj.get("name", ""), proj.get("technologies", "")
            label = name + (f" — {tech}" if tech else "")
            if label.strip():
                body_line(label, bold=True)
            if proj.get("description"):
                body_line(proj["description"])
            pdf.ln(1)
        pdf.ln(1)

    def render_certifications():
        certs = content.get("certifications") or []
        if not certs:
            return
        section(labels["certifications"])
        emphasize = config.get("emphasize_certifications", False)
        for cert in certs:
            name, issuer, date = cert.get("name", ""), cert.get("issuer", ""), cert.get("date", "")
            line = name + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                body_line(line, bold=emphasize, fill=emphasize)
        pdf.ln(2)

    def render_achievements():
        if not content.get("achievements"):
            return
        section(labels["achievements"])
        for a in content["achievements"]:
            if str(a).strip():
                body_line(f"  -  {a}")
        pdf.ln(2)

    def render_languages():
        strs = _language_strings(content)
        if not strs:
            return
        section(labels["languages"])
        body_line("  ·  ".join(strs))
        pdf.ln(2)

    def render_interests():
        items = [str(i) for i in (content.get("interests") or []) if str(i).strip()]
        if not items:
            return
        section(labels["interests"])
        body_line("  ·  ".join(items))

    def render_custom_sections():
        # Unlike every other section, there's no single fixed label -- each
        # entry's heading is whatever the candidate actually typed as its
        # title (never hardcoded, never defaulted to a specific example).
        for cs in (content.get("customSections") or []):
            cs_title = (cs.get("title") or "").strip()
            cs_body = (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            section(cs_title or "Additional Information")
            if cs_body:
                body_line(cs_body)
            pdf.ln(1)

    renderers = {
        "summary": render_summary, "skills": render_skills, "experience": render_experience,
        "education": render_education, "projects": render_projects, "certifications": render_certifications,
        "achievements": render_achievements, "languages": render_languages, "interests": render_interests,
        "customSections": render_custom_sections,
    }
    for key in order:
        if key in sections:
            renderers[key]()

    return bytes(pdf.output())


def _build_docx_single_column(content: dict, title: str, sections: list[str], *, template_id: str,
                               style: dict | None = None) -> bytes:
    """`style` (see _resolve_style) carries the candidate's own explicit
    font/spacing override on top of this template's own design; absent/None
    means "use the template's own design, no scaling" — today's exact
    behavior, byte-for-byte."""
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from services.skill_categories import group_skills_by_category

    config = SINGLE_COLUMN_CONFIGS[template_id]
    accent_rgb = RGBColor(*config["accent"])
    style = style or {}
    docx_font = style.get("docx_font") or config["docx_font"]
    font_scale = style.get("font_scale", 1.0)
    spacing_scale = style.get("spacing_scale", 1.0)
    Pt = _docx_pt_scaler(font_scale)
    PtSpace = _docx_space_scaler(spacing_scale)
    labels = config["labels"]
    order_spec = config["order"]
    order = order_spec(bool(content.get("experience"))) if callable(order_spec) else order_spec

    pi = content.get("personalInfo", {})
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(36)
        sec.left_margin = sec.right_margin = Pt(54)
    doc.styles["Normal"].paragraph_format.line_spacing = spacing_scale

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = name_para.add_run(pi.get("fullName", title))
    run.font.name = docx_font
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = accent_rgb

    if pi.get("jobTitle"):
        jt = doc.add_paragraph()
        r = jt.add_run(pi["jobTitle"])
        r.font.name = docx_font
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    contact_parts = [s for s in [pi.get("location"), pi.get("phone"), pi.get("email"),
                                    pi.get("linkedin"), pi.get("github"), pi.get("website")] if s]
    if contact_parts:
        cp = doc.add_paragraph("  |  ".join(contact_parts))
        cp.runs[0].font.name = docx_font
        cp.runs[0].font.size = Pt(9)
        cp.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    def add_section(heading: str):
        h = doc.add_paragraph(heading.upper())
        run = h.runs[0]
        run.font.name = docx_font
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = accent_rgb
        h.paragraph_format.space_before = PtSpace(10)
        # (python-docx has no simple paragraph-border API; the accent color +
        # bold + spacing is the heading treatment here, same as the classic builder)

    def add_paragraph(text: str, *, bold: bool = False, style: str | None = None):
        p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
        for r in p.runs:
            r.font.name = docx_font
            if bold:
                r.font.bold = True
        return p

    def render_summary():
        if content.get("summary"):
            add_section(labels["summary"]); add_paragraph(content["summary"])

    def render_skills():
        names = _skill_names(content)
        if not names:
            return
        add_section(labels["skills"])
        if config.get("group_skills"):
            for cat, cat_skills in group_skills_by_category(names):
                add_paragraph(f"{cat}: {', '.join(cat_skills)}")
        else:
            add_paragraph("  •  ".join(names))

    def render_experience():
        if not content.get("experience"):
            return
        add_section(labels["experience"])
        for exp in content["experience"]:
            p = doc.add_paragraph()
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            dates = f"  ({exp.get('startDate', '')} - {end})" if exp.get("startDate") else ""
            r = p.add_run(f"{exp.get('position', '')} — {exp.get('company', '')}{dates}")
            r.font.name = docx_font
            r.font.bold = True
            for bullet in exp.get("bullets", []):
                if bullet.strip():
                    bp = doc.add_paragraph(bullet, style="List Bullet")
                    bp.paragraph_format.left_indent = Pt(18)
                    for r2 in bp.runs:
                        r2.font.name = docx_font

    def render_education():
        if not content.get("education"):
            return
        add_section(labels["education"])
        for edu in content["education"]:
            p = doc.add_paragraph()
            deg = ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            end_date = edu.get("endDate", "")
            r = p.add_run((deg or edu.get("institution", "")) + (f"  ({end_date})" if end_date else ""))
            r.font.name = docx_font
            r.font.bold = True
            if deg and edu.get("institution"):
                add_paragraph(edu["institution"])

    def render_projects():
        if not content.get("projects"):
            return
        add_section(labels["projects"])
        for proj in content["projects"]:
            name, tech = proj.get("name", ""), proj.get("technologies", "")
            if name or tech:
                add_paragraph(name + (f"  ({tech})" if tech else ""), bold=True)
            if proj.get("description"):
                add_paragraph(proj["description"])

    def render_certifications():
        certs = content.get("certifications") or []
        if not certs:
            return
        add_section(labels["certifications"])
        emphasize = config.get("emphasize_certifications", False)
        for cert in certs:
            name, issuer, date = cert.get("name", ""), cert.get("issuer", ""), cert.get("date", "")
            line = name + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                add_paragraph(line, bold=emphasize)

    def render_achievements():
        if not content.get("achievements"):
            return
        add_section(labels["achievements"])
        for a in content["achievements"]:
            if str(a).strip():
                add_paragraph(str(a), style="List Bullet")

    def render_languages():
        strs = _language_strings(content)
        if not strs:
            return
        add_section(labels["languages"])
        add_paragraph("  •  ".join(strs))

    def render_interests():
        items = [str(i) for i in (content.get("interests") or []) if str(i).strip()]
        if not items:
            return
        add_section(labels["interests"])
        add_paragraph("  •  ".join(items))

    def render_custom_sections():
        for cs in (content.get("customSections") or []):
            cs_title = (cs.get("title") or "").strip()
            cs_body = (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            add_section(cs_title or "Additional Information")
            if cs_body:
                add_paragraph(cs_body)

    renderers = {
        "summary": render_summary, "skills": render_skills, "experience": render_experience,
        "education": render_education, "projects": render_projects, "certifications": render_certifications,
        "achievements": render_achievements, "languages": render_languages, "interests": render_interests,
        "customSections": render_custom_sections,
    }
    for key in order:
        if key in sections:
            renderers[key]()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── "Classic" family, template-aware ────────────────────────────────────────
# modern/professional/minimal/creative/executive each get their OWN look here,
# driven entirely by TEMPLATE_SPECS[template_id]'s accent/font/layoutFamily —
# the exact same fields frontend/components/ResumeTemplates.tsx's template
# picker reads (shared/template-specs.json is the single configuration layer;
# see module docstring). Before this, all five mapped to one generic,
# template-blind design (_build_pdf/_build_docx below, kept only as the
# "classic-legacy" fallback for an unrecognized layoutFamily and for direct
# unit-test callers that don't go through the registry).

def _classic_section_renderer(pdf, font: str, safe, accent: tuple[int, int, int], *, style: str):
    """Returns (section, body_line) closures for single-column classic
    layouts, so Professional/Minimal (and the legacy default) can each get
    their own heading treatment without duplicating the body-line logic.
    `style`: "boxed" (filled heading bar, e.g. legacy), "bordered-caps"
    (centered-serif look, e.g. Professional), or "tracked-light" (thin
    uppercase gray heading, e.g. Minimal)."""
    def section(heading: str):
        if style == "tracked-light":
            pdf.ln(2)
            pdf.set_font(font, "", 9)
            pdf.set_text_color(*accent)
            pdf.set_x(pdf.l_margin)
            pdf.cell(0, 6, safe(heading.upper()), new_x="LMARGIN", new_y="NEXT")
        elif style == "bordered-caps":
            pdf.ln(2)
            pdf.set_font(font, "B", 10)
            pdf.set_text_color(*accent)
            pdf.set_x(pdf.l_margin)
            pdf.cell(0, 6, safe(heading.upper()), new_x="LMARGIN", new_y="NEXT")
            pdf.set_draw_color(*accent)
            pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
            pdf.ln(1.5)
        else:  # "boxed" — legacy look, kept for the fallback path
            pdf.set_font(font, "B", 10)
            pdf.set_text_color(*accent)
            pdf.set_fill_color(min(accent[0] + 209, 255), min(accent[1] + 182, 255), min(accent[2] + 80, 255))
            pdf.set_x(pdf.l_margin)
            pdf.cell(0, 7, safe(heading.upper()), new_x="LMARGIN", new_y="NEXT", fill=True)
        pdf.set_text_color(30, 30, 30)
        pdf.set_font(font, "", 10)

    def body_line(text: str, size: int = 10, bold: bool = False, grey: bool = False, width: float = 0):
        pdf.set_x(pdf.l_margin)
        pdf.set_font(font, "B" if bold else "", size)
        pdf.set_text_color(120, 120, 120) if grey else pdf.set_text_color(30, 30, 30)
        pdf.multi_cell(width, 5 if not bold else 6, safe(text))
        pdf.set_text_color(30, 30, 30)

    return section, body_line


def _classic_pdf_single_column(content: dict, sections: list[str], *, accent: tuple[int, int, int],
                                font_family: str, centered_header: bool, heading_style: str,
                                font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Professional (centered, serif, bordered-caps headings) and Minimal
    (left, sans, thin tracked-light headings) both fit this one shape —
    only the header layout and heading treatment differ, exactly like their
    frontend components (ProfessionalTemplate.tsx / MinimalTemplate.tsx).
    font_scale/spacing_scale come from the candidate's own Font-Size/Spacing
    picker choice (see _resolve_style) — 1.0/1.0 is today's exact output."""
    from fpdf import FPDF

    pi = content.get("personalInfo", {})
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    font, safe = _pdf_font(pdf, font_family)
    _apply_pdf_scale(pdf, font_scale, spacing_scale)

    name = safe(pi.get("fullName") or content.get("title") or "")
    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location")] if s]
    if centered_header:
        pdf.set_font(font, "B", 20)
        pdf.set_text_color(20, 20, 20)
        pdf.cell(0, 10, name, align="C", new_x="LMARGIN", new_y="NEXT")
        if pi.get("jobTitle"):
            pdf.set_font(font, "", 11)
            pdf.set_text_color(90, 90, 90)
            pdf.cell(0, 6, safe(pi["jobTitle"]), align="C", new_x="LMARGIN", new_y="NEXT")
        if contact_parts:
            pdf.set_font(font, "", 9)
            pdf.set_text_color(110, 110, 110)
            pdf.cell(0, 6, safe("  |  ".join(contact_parts)), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.set_draw_color(*accent)
        pdf.set_line_width(0.6)
        pdf.line(pdf.l_margin, pdf.get_y() + 1, pdf.w - pdf.r_margin, pdf.get_y() + 1)
        pdf.ln(5)
    else:
        pdf.set_font(font, "", 22)
        pdf.set_text_color(20, 20, 20)
        pdf.cell(0, 11, name, new_x="LMARGIN", new_y="NEXT")
        if pi.get("jobTitle"):
            pdf.set_font(font, "", 10)
            pdf.set_text_color(*accent)
            pdf.cell(0, 6, safe(pi["jobTitle"]), new_x="LMARGIN", new_y="NEXT")
        if contact_parts:
            pdf.set_font(font, "", 9)
            pdf.set_text_color(140, 140, 140)
            pdf.cell(0, 6, safe("   ".join(contact_parts)), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    section, body_line = _classic_section_renderer(pdf, font, safe, accent, style=heading_style)

    if "summary" in sections and content.get("summary"):
        if heading_style == "tracked-light":
            # Minimal renders the summary as a plain left-bordered quote, no heading.
            pdf.set_draw_color(220, 220, 220)
            pdf.set_line_width(0.5)
            x0, y0 = pdf.l_margin, pdf.get_y()
            pdf.set_x(x0 + 3)
            body_line(content["summary"], grey=True, width=pdf.w - pdf.r_margin - x0 - 3)
            pdf.line(x0, y0, x0, pdf.get_y())
            pdf.ln(3)
        else:
            section("Professional Summary")
            body_line(content["summary"])
            pdf.ln(3)

    if "experience" in sections and content.get("experience"):
        section("Experience")
        for exp in content["experience"]:
            position = exp.get("position", "")
            company = exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            date_str = f"  ({start} - {end})" if start else ""
            label = position + (f" — {company}" if company else "") + date_str
            body_line(label, bold=True)
            for bullet in exp.get("bullets", []):
                if str(bullet).strip():
                    body_line(f"  -  {bullet}")
            pdf.ln(2)

    if "projects" in sections and content.get("projects"):
        section("Projects")
        for proj in content["projects"]:
            name_ = proj.get("name", "")
            tech = proj.get("technologies", "")
            label = name_ + (f"  ({tech})" if tech else "")
            if label.strip():
                body_line(label, bold=True)
            if proj.get("description"):
                body_line(proj["description"])
            pdf.ln(1)

    if "education" in sections and content.get("education"):
        section("Education")
        for edu in content["education"]:
            inst = edu.get("institution", "")
            deg = ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            end_date = edu.get("endDate", "")
            body_line(inst + (f"  ({end_date})" if end_date else ""), bold=True)
            if deg:
                body_line(deg)
            pdf.ln(2)

    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            section("Skills")
            body_line("  -  ".join(skill_names))
            pdf.ln(2)

    if "certifications" in sections and content.get("certifications"):
        section("Certifications")
        for cert in content["certifications"]:
            name_ = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            line = name_ + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                body_line(line)
        pdf.ln(2)

    if "achievements" in sections and content.get("achievements"):
        section("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                body_line(f"  -  {a}")
        pdf.ln(2)

    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            section("Languages")
            body_line("  -  ".join(lang_strs))
            pdf.ln(2)

    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            section("Interests")
            body_line("  -  ".join(interests))
            pdf.ln(2)

    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title = (cs.get("title") or "").strip()
            cs_body = (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            section(cs_title or "Additional Information")
            if cs_body:
                body_line(cs_body)
            pdf.ln(1)

    return bytes(pdf.output())


def _classic_pdf_sidebar(content: dict, sections: list[str], *, accent: tuple[int, int, int], font_family: str,
                          font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Modern: colored left sidebar (contact/skills/languages) + white main
    column (summary/experience/projects/education/customSections) — mirrors
    ModernTemplate.tsx's 38%-width dark sidebar. The sidebar band is drawn
    via an FPDF subclass's header() override so it repeats on every page
    the main column overflows onto; sidebar CONTENT (contact/skills/
    languages) is only rendered once, on page 1, since it's short by nature
    and a resume's contact info repeating per page would look wrong anyway.
    ModernTemplate.tsx itself doesn't render achievements/certifications/
    interests at all -- rather than silently drop real candidate data the
    editor collected, this PDF still includes them (appended to the main
    column) if present. See module docstring for the multi-page caveat."""
    from fpdf import FPDF

    pi = content.get("personalInfo", {})
    sidebar_w = 68.0  # mm — ≈ the frontend's 38% column at A4's 210mm width
    pad = 6.0

    class _SidebarPDF(FPDF):
        def header(self):
            self.set_fill_color(*accent)
            self.rect(0, 0, sidebar_w, self.h, style="F")

    pdf = _SidebarPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    font, safe = _pdf_font(pdf, font_family)
    _apply_pdf_scale(pdf, font_scale, spacing_scale)

    # ── Sidebar content (page 1 only — see docstring) ──
    sb_x, sb_w = pad, sidebar_w - 2 * pad
    pdf.set_xy(sb_x, 14)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font(font, "B", 14)
    pdf.set_x(sb_x)
    pdf.multi_cell(sb_w, 6, safe(pi.get("fullName") or content.get("title") or ""), align="C")
    if pi.get("jobTitle"):
        pdf.set_font(font, "", 9)
        pdf.set_text_color(200, 210, 240)
        pdf.set_x(sb_x)
        pdf.multi_cell(sb_w, 5, safe(pi["jobTitle"]), align="C")
    pdf.ln(3)

    def sb_heading(text: str):
        pdf.set_x(sb_x)
        pdf.set_font(font, "B", 8)
        pdf.set_text_color(190, 205, 240)
        pdf.multi_cell(sb_w, 5, safe(text.upper()))

    def sb_line(text: str, size: float = 8.5):
        pdf.set_x(sb_x)
        pdf.set_font(font, "", size)
        pdf.set_text_color(230, 235, 250)
        pdf.multi_cell(sb_w, 4.6, safe(text))

    contact_lines = [v for v in [pi.get("location"), pi.get("phone"), pi.get("email"), pi.get("linkedin"), pi.get("github")] if v]
    if contact_lines:
        sb_heading("Contact")
        for c in contact_lines:
            sb_line(c)
        pdf.ln(3)

    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            sb_heading("Skills")
            for s in skill_names:
                sb_line(f"- {s}")
            pdf.ln(3)

    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            sb_heading("Languages")
            for l in lang_strs:
                sb_line(l)

    # ── Main column ──
    main_x = sidebar_w + pad
    pdf.set_left_margin(main_x)
    pdf.set_right_margin(14)
    pdf.set_xy(main_x, 14)
    section, body_line = _classic_section_renderer(pdf, font, safe, accent, style="bordered-caps")
    pdf.ln(-2)  # bordered-caps opens with ln(2); the very first heading shouldn't add top padding

    if "summary" in sections and content.get("summary"):
        section("Professional Summary")
        body_line(content["summary"])
        pdf.ln(2)

    if "experience" in sections and content.get("experience"):
        section("Work Experience")
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            date_str = f"  ({start} - {end})" if start else ""
            body_line(position + (f" - {company}" if company else "") + date_str, bold=True)
            for bullet in exp.get("bullets", []):
                if str(bullet).strip():
                    body_line(f"  -  {bullet}")
            pdf.ln(2)

    if "projects" in sections and content.get("projects"):
        section("Projects")
        for proj in content["projects"]:
            name_, tech = proj.get("name", ""), proj.get("technologies", "")
            label = name_ + (f"  ({tech})" if tech else "")
            if label.strip():
                body_line(label, bold=True)
            if proj.get("description"):
                body_line(proj["description"])
            pdf.ln(1)

    if "education" in sections and content.get("education"):
        section("Education")
        for edu in content["education"]:
            inst = edu.get("institution", "")
            deg = ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            end_date = edu.get("endDate", "")
            body_line(inst + (f"  ({end_date})" if end_date else ""), bold=True)
            if deg:
                body_line(deg)
            pdf.ln(2)

    # Not rendered by ModernTemplate.tsx's Preview at all -- appended here so
    # real candidate data is never silently dropped (see docstring).
    if "certifications" in sections and content.get("certifications"):
        section("Certifications")
        for cert in content["certifications"]:
            name_, issuer, date = cert.get("name", ""), cert.get("issuer", ""), cert.get("date", "")
            line = name_ + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                body_line(line)
        pdf.ln(2)

    if "achievements" in sections and content.get("achievements"):
        section("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                body_line(f"  -  {a}")
        pdf.ln(2)

    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            section("Interests")
            body_line("  -  ".join(interests))
            pdf.ln(2)

    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title = (cs.get("title") or "").strip()
            cs_body = (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            section(cs_title or "Additional Information")
            if cs_body:
                body_line(cs_body)
            pdf.ln(1)

    return bytes(pdf.output())


def _classic_pdf_sidebar_dark(content: dict, sections: list[str], *, accent: tuple[int, int, int], font_family: str,
                               font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Executive: a dark full-width header band, then a 2-column grid below
    it — main column (summary/experience/projects/achievements/
    customSections) beside a light-gray sidebar (education/skills/
    certifications) — mirroring ExecutiveTemplate.tsx's `grid-cols-[1fr,0.45fr]`.
    The grid is rendered as two independently-flowing column passes starting
    at the same Y; like _classic_pdf_sidebar, correctness is single-page-first
    — see module docstring for the multi-page caveat."""
    from fpdf import FPDF

    pi = content.get("personalInfo", {})
    pdf = FPDF()
    pdf.set_auto_page_break(auto=False)
    pdf.add_page()
    pdf.set_margins(0, 0, 0)
    font, safe = _pdf_font(pdf, font_family)
    _apply_pdf_scale(pdf, font_scale, spacing_scale)
    page_w = pdf.w

    # ── Dark header band ──
    band_h = 30.0
    pdf.set_fill_color(*accent)
    pdf.rect(0, 0, page_w, band_h, style="F")
    pdf.set_xy(16, 8)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font(font, "B", 17)
    pdf.cell(page_w - 32, 8, safe((pi.get("fullName") or content.get("title") or "").upper()), new_x="LMARGIN", new_y="NEXT")
    if pi.get("jobTitle"):
        pdf.set_x(16)
        pdf.set_font(font, "", 10)
        pdf.set_text_color(200, 200, 205)
        pdf.cell(page_w - 32, 6, safe(pi["jobTitle"].upper()), new_x="LMARGIN", new_y="NEXT")
    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location"), pi.get("linkedin")] if s]
    if contact_parts:
        pdf.set_x(16)
        pdf.set_font(font, "", 8.5)
        pdf.set_text_color(180, 180, 186)
        pdf.cell(page_w - 32, 5, safe("    ".join(contact_parts)), new_x="LMARGIN", new_y="NEXT")

    # ── 2-column grid below the band ──
    grid_top = band_h + 8
    main_x, main_w = 16.0, page_w * 0.63 - 20
    side_x, side_w = page_w * 0.65, page_w * 0.35 - 16 - 4
    divider_x = page_w * 0.645

    def divider_heading(pdf_, x, w, text):
        pdf_.set_xy(x, pdf_.get_y() + 2)
        pdf_.set_font(font, "B", 8.5)
        pdf_.set_text_color(*accent)
        pdf_.cell(w, 5, safe(text.upper()), align="C", new_x="LMARGIN", new_y="NEXT")
        pdf_.set_draw_color(210, 210, 214)
        pdf_.line(x, pdf_.get_y() + 0.5, x + w, pdf_.get_y() + 0.5)
        pdf_.ln(2.5)
        pdf_.set_x(x)

    def col_body(pdf_, x, w, text, bold=False, size=9.5, grey=False):
        pdf_.set_x(x)
        pdf_.set_font(font, "B" if bold else "", size)
        pdf_.set_text_color(120, 120, 120) if grey else pdf_.set_text_color(30, 30, 30)
        pdf_.multi_cell(w, 4.6, safe(text))
        pdf_.set_text_color(30, 30, 30)
        pdf_.set_x(x)

    # Main column
    pdf.set_xy(main_x, grid_top)
    if "summary" in sections and content.get("summary"):
        divider_heading(pdf, main_x, main_w, "Executive Profile")
        col_body(pdf, main_x, main_w, content["summary"])
        pdf.ln(2)
    if "experience" in sections and content.get("experience"):
        divider_heading(pdf, main_x, main_w, "Career History")
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            col_body(pdf, main_x, main_w, position, bold=True)
            meta = company + (f"  ({start} - {end})" if start else "")
            if meta.strip():
                col_body(pdf, main_x, main_w, meta, grey=True, size=8.5)
            for b in exp.get("bullets", []):
                if str(b).strip():
                    col_body(pdf, main_x, main_w, f"  -  {b}")
            pdf.ln(1.5)
    if "projects" in sections and content.get("projects"):
        divider_heading(pdf, main_x, main_w, "Projects")
        for proj in content["projects"]:
            name_, tech = proj.get("name", ""), proj.get("technologies", "")
            label = name_ + (f"  ({tech})" if tech else "")
            if label.strip():
                col_body(pdf, main_x, main_w, label, bold=True)
            if proj.get("description"):
                col_body(pdf, main_x, main_w, proj["description"])
            pdf.ln(1)
    if "achievements" in sections and content.get("achievements"):
        divider_heading(pdf, main_x, main_w, "Key Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                col_body(pdf, main_x, main_w, f"  -  {a}")
        pdf.ln(1.5)
    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title, cs_body = (cs.get("title") or "").strip(), (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            divider_heading(pdf, main_x, main_w, cs_title or "Additional Information")
            if cs_body:
                col_body(pdf, main_x, main_w, cs_body)
            pdf.ln(1)
    main_end_y = pdf.get_y()

    # Sidebar column (independent flow, same starting Y)
    pdf.set_fill_color(248, 249, 250)
    pdf.rect(page_w * 0.645, grid_top - 2, page_w * 0.355, 297 - grid_top, style="F")
    pdf.set_draw_color(*accent)
    pdf.line(divider_x, grid_top - 2, divider_x, 290)

    def side_heading(text: str):
        pdf.set_x(side_x)
        pdf.set_font(font, "B", 8)
        pdf.set_text_color(90, 90, 90)
        pdf.cell(side_w, 5, safe(text.upper()), new_x="LMARGIN", new_y="NEXT")
        pdf.set_x(side_x)
        pdf.ln(1)

    pdf.set_xy(side_x, grid_top)
    if "education" in sections and content.get("education"):
        side_heading("Education")
        for edu in content["education"]:
            col_body(pdf, side_x, side_w, edu.get("degree", ""), bold=True, size=9)
            if edu.get("field"):
                col_body(pdf, side_x, side_w, edu["field"], grey=True, size=8.5)
            if edu.get("institution"):
                col_body(pdf, side_x, side_w, edu["institution"], grey=True, size=8.5)
            pdf.ln(1.5)
    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            side_heading("Core Competencies")
            for s in skill_names:
                col_body(pdf, side_x, side_w, f"- {s}", size=8.5)
            pdf.ln(1.5)
    if "certifications" in sections and content.get("certifications"):
        side_heading("Certifications")
        for cert in content["certifications"]:
            col_body(pdf, side_x, side_w, cert.get("name", ""), bold=True, size=8.5)
            meta = " · ".join(filter(None, [cert.get("issuer"), cert.get("date")]))
            if meta:
                col_body(pdf, side_x, side_w, meta, grey=True, size=8)
            pdf.ln(1)
    # Not rendered by ExecutiveTemplate.tsx's Preview at all -- appended here
    # so real candidate data is never silently dropped (see Modern's docstring).
    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            side_heading("Languages")
            for l in lang_strs:
                col_body(pdf, side_x, side_w, l, size=8.5)
            pdf.ln(1.5)
    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            side_heading("Interests")
            col_body(pdf, side_x, side_w, "  -  ".join(interests), size=8.5)
    side_end_y = pdf.get_y()

    pdf.set_y(max(main_end_y, side_end_y) + 4)
    return bytes(pdf.output())


def _classic_pdf_header_band(content: dict, sections: list[str], *, accent: tuple[int, int, int], font_family: str,
                              font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Creative: a full-width gradient header (approximated with thin
    interpolated color strips — fpdf2 has no native gradient fill), a thin
    gradient accent bar down the left edge, left-bordered experience/project
    entries, and a 2-column Education | Skills grid near the end — mirroring
    CreativeTemplate.tsx. `accent` is the template's purple; blue is mixed in
    to reproduce the purple→blue gradient CreativeTemplate.tsx uses."""
    from fpdf import FPDF

    pi = content.get("personalInfo", {})
    blue = (37, 99, 235)  # CreativeTemplate.tsx's gradient end (blue-600)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(0, 0, 0)
    font, safe = _pdf_font(pdf, font_family)
    _apply_pdf_scale(pdf, font_scale, spacing_scale)
    page_w = pdf.w

    # ── Gradient header band ──
    band_h = 26.0
    strips = 60
    for i in range(strips):
        t = i / (strips - 1)
        pdf.set_fill_color(*_mix_rgb(accent, blue, t))
        pdf.rect(page_w * i / strips, 0, page_w / strips + 0.5, band_h, style="F")
    pdf.set_xy(16, 6)
    pdf.set_text_color(255, 255, 255)
    pdf.set_font(font, "B", 15)
    pdf.cell(page_w - 32, 7, safe(pi.get("fullName") or content.get("title") or ""), new_x="LMARGIN", new_y="NEXT")
    if pi.get("jobTitle"):
        pdf.set_x(16)
        pdf.set_font(font, "", 9.5)
        pdf.set_text_color(230, 225, 250)
        pdf.cell(page_w - 32, 5, safe(pi["jobTitle"]), new_x="LMARGIN", new_y="NEXT")
    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location")] if s]
    if contact_parts:
        pdf.set_x(16)
        pdf.set_font(font, "", 8.5)
        pdf.set_text_color(225, 220, 250)
        pdf.cell(page_w - 32, 5, safe("    ".join(contact_parts)), new_x="LMARGIN", new_y="NEXT")

    # ── Thin left accent bar + main content margin ──
    content_x = 14.0
    pdf.set_left_margin(content_x)
    pdf.set_right_margin(12)
    pdf.set_xy(content_x, band_h + 6)

    def section(heading: str):
        pdf.set_x(content_x)
        pdf.set_font(font, "B", 8.5)
        pdf.set_text_color(*accent)
        pdf.cell(0, 5, safe(heading.upper()), new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(30, 30, 30)
        pdf.set_font(font, "", 10)
        pdf.set_x(content_x)

    def bordered_body(text: str, bold: bool = False, size: float = 9.5, color: tuple[int, int, int] | None = None,
                       width: float | None = None):
        """A left-border-accented block, like CreativeTemplate.tsx's
        `border-l-2 border-purple-100` experience/project entries."""
        w = width if width is not None else (page_w - content_x - pdf.r_margin - 4)
        x0, y0 = content_x, pdf.get_y()
        pdf.set_x(content_x + 4)
        pdf.set_font(font, "B" if bold else "", size)
        pdf.set_text_color(*(color or (30, 30, 30)))
        pdf.multi_cell(w, 4.8, safe(text))
        pdf.set_text_color(30, 30, 30)
        pdf.set_draw_color(*_mix_rgb(accent, (255, 255, 255), 0.75))
        pdf.set_line_width(0.6)
        pdf.line(x0, y0, x0, pdf.get_y())
        pdf.set_x(content_x)

    if "summary" in sections and content.get("summary"):
        section("About Me")
        pdf.set_x(content_x)
        pdf.multi_cell(0, 4.8, safe(content["summary"]))
        pdf.ln(2)

    if "experience" in sections and content.get("experience"):
        section("Experience")
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            bordered_body(position, bold=True)
            meta = company + (f"  ({start} - {end})" if start else "")
            if meta.strip():
                bordered_body(meta, color=_mix_rgb(accent, blue, 0.5))
            for b in exp.get("bullets", []):
                if str(b).strip():
                    bordered_body(f"›  {b}")
            pdf.ln(1.5)

    if "projects" in sections and content.get("projects"):
        section("Projects")
        for proj in content["projects"]:
            name_, tech = proj.get("name", ""), proj.get("technologies", "")
            bordered_body(name_, bold=True)
            if tech:
                bordered_body(tech, color=_mix_rgb(accent, blue, 0.5))
            if proj.get("description"):
                bordered_body(proj["description"])
            pdf.ln(1)

    # 2-column grid: Education | Skills (with bars) — CreativeTemplate.tsx's
    # `grid grid-cols-2` block near the end.
    has_edu = "education" in sections and content.get("education")
    skill_names = _skill_names(content) if "skills" in sections else []
    if has_edu or skill_names:
        col_w = (page_w - content_x - pdf.r_margin - 6) / 2
        left_x, right_x = content_x, content_x + col_w + 6
        grid_top = pdf.get_y() + 1
        if has_edu:
            pdf.set_xy(left_x, grid_top)
            pdf.set_font(font, "B", 8.5)
            pdf.set_text_color(*accent)
            pdf.cell(col_w, 5, "EDUCATION", new_x="LMARGIN", new_y="NEXT")
            for edu in content["education"]:
                pdf.set_x(left_x)
                pdf.set_font(font, "B", 9)
                pdf.set_text_color(30, 30, 30)
                pdf.multi_cell(col_w, 4.5, safe(edu.get("degree", "")))
                if edu.get("institution"):
                    pdf.set_x(left_x)
                    pdf.set_font(font, "", 8.5)
                    pdf.set_text_color(120, 120, 120)
                    pdf.multi_cell(col_w, 4.3, safe(edu["institution"]))
                pdf.ln(1)
        edu_end_y = pdf.get_y()
        if skill_names:
            pdf.set_xy(right_x, grid_top)
            pdf.set_font(font, "B", 8.5)
            pdf.set_text_color(*accent)
            pdf.cell(col_w, 5, "SKILLS", new_x="LMARGIN", new_y="NEXT")
            for s in skill_names[:8]:
                pdf.set_x(right_x)
                pdf.set_font(font, "", 8.5)
                pdf.set_text_color(60, 60, 60)
                pdf.multi_cell(col_w, 4.3, safe(s))
                # a filled bar approximating the frontend's gradient skill-level bar
                bar_y = pdf.get_y() + 0.3
                pdf.set_fill_color(230, 230, 235)
                pdf.rect(right_x, bar_y, col_w, 1.3, style="F")
                pdf.set_fill_color(*_mix_rgb(accent, blue, 0.5))
                pdf.rect(right_x, bar_y, col_w * 0.75, 1.3, style="F")
                pdf.set_y(bar_y + 3)
        skills_end_y = pdf.get_y()
        pdf.set_y(max(edu_end_y, skills_end_y) + 2)
        pdf.set_x(content_x)

    if "certifications" in sections and content.get("certifications"):
        section("Certifications")
        for cert in content["certifications"]:
            name_, issuer, date = cert.get("name", ""), cert.get("issuer", ""), cert.get("date", "")
            line = name_ + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                pdf.set_x(content_x)
                pdf.multi_cell(0, 4.6, safe(line))
        pdf.ln(2)

    if "achievements" in sections and content.get("achievements"):
        section("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                pdf.set_x(content_x)
                pdf.multi_cell(0, 4.6, safe(f"›  {a}"))
        pdf.ln(2)

    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            section("Languages")
            pdf.set_x(content_x)
            pdf.multi_cell(0, 4.6, safe("  -  ".join(lang_strs)))
            pdf.ln(2)

    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            section("Interests")
            pdf.set_x(content_x)
            pdf.multi_cell(0, 4.6, safe("  -  ".join(interests)))
            pdf.ln(2)

    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title, cs_body = (cs.get("title") or "").strip(), (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            section(cs_title or "Additional Information")
            if cs_body:
                pdf.set_x(content_x)
                pdf.multi_cell(0, 4.6, safe(cs_body))
            pdf.ln(1)

    return bytes(pdf.output())


def _build_pdf_classic(content: dict, title: str, sections: list[str], *, template_id: str,
                        style: dict | None = None) -> bytes:
    """Dispatches to the classic-family layout matching template_id's own
    spec (shared/template-specs.json) — the single configuration layer both
    the frontend template picker and this exporter read. `style` (see
    _resolve_style) carries the candidate's own explicit font/spacing
    override on top of that; absent/None means "use the template's own
    design, no scaling" — today's exact behavior. Falls back to the
    template-blind legacy design only for an id whose layoutFamily this
    function doesn't recognize (should not happen for a valid registry id)."""
    spec = TEMPLATE_SPECS.get(template_id, {})
    style = style or {}
    accent = _hex_to_rgb(spec.get("accent") or "#1e40af")
    font_family = style.get("font_family") or spec.get("font") or "sans"
    font_scale = style.get("font_scale", 1.0)
    spacing_scale = style.get("spacing_scale", 1.0)
    layout = spec.get("layoutFamily") or ""

    if layout == "sidebar-left":
        return _classic_pdf_sidebar(content, sections, accent=accent, font_family=font_family,
                                     font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "sidebar-dark":
        return _classic_pdf_sidebar_dark(content, sections, accent=accent, font_family=font_family,
                                          font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "header-band":
        return _classic_pdf_header_band(content, sections, accent=accent, font_family=font_family,
                                         font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "single-column-serif":
        return _classic_pdf_single_column(content, sections, accent=accent, font_family=font_family,
                                           centered_header=True, heading_style="bordered-caps",
                                           font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "single-column-minimal":
        return _classic_pdf_single_column(content, sections, accent=accent, font_family=font_family,
                                           centered_header=False, heading_style="tracked-light",
                                           font_scale=font_scale, spacing_scale=spacing_scale)
    return _build_pdf(content, title, sections)


# ── "Classic" family DOCX, template-aware ───────────────────────────────────
# DOCX counterpart of the PDF layouts above — same accent/font/layoutFamily
# read from TEMPLATE_SPECS. Word's table/paragraph flow auto-paginates, so
# unlike the PDF versions, multi-column layouts here don't have a
# single-page-first caveat. What DOCX genuinely can't reproduce: the
# gradient header (approximated as one solid accent color instead of a
# blend) and per-paragraph left-accent borders on Experience/Project entries
# (approximated with accent-colored bold sub-text instead) — see each
# function's docstring for the specific trade-off.

def _docx_shade_cell(cell, hex_color: str):
    """Solid background fill for a table cell. python-docx has no high-level
    API for this — direct oxml is the standard, documented workaround."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def _docx_bottom_border(paragraph, rgb: tuple[int, int, int]):
    """A colored rule under a paragraph (e.g. under a heading or the header
    block) — approximates the frontend's `border-b` accent divider."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "%02X%02X%02X" % rgb)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_no_borders(table):
    """Removes a python-docx table's default visible grid — layout tables
    (used here purely for column positioning, e.g. the sidebar) must stay
    invisible, unlike a real data table."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tbl_pr.append(borders)


def _classic_docx_single_column(content: dict, sections: list[str], *, accent: tuple[int, int, int],
                                 docx_font: str, centered_header: bool, heading_style: str,
                                 font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Professional (centered, serif, bordered-caps) and Minimal (left, sans,
    thin tracked-light headings) — DOCX counterpart of _classic_pdf_single_column.
    font_scale/spacing_scale come from the candidate's own Font-Size/Spacing
    picker choice (see _resolve_style) — 1.0/1.0 is today's exact output."""
    from docx import Document
    from docx.shared import RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    Pt = _docx_pt_scaler(font_scale)
    PtSpace = _docx_space_scaler(spacing_scale)

    pi = content.get("personalInfo", {})
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(36)
        sec.left_margin = sec.right_margin = Pt(54)
    doc.styles["Normal"].font.name = docx_font
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.line_spacing = spacing_scale
    accent_rgb = RGBColor(*accent)

    name_para = doc.add_paragraph()
    if centered_header:
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name_para.add_run(pi.get("fullName") or content.get("title") or "")
    r.font.size, r.font.bold, r.font.name = Pt(22), True, docx_font
    r.font.color.rgb = RGBColor(0x14, 0x14, 0x14)

    last_header_para = name_para
    if pi.get("jobTitle"):
        jt = doc.add_paragraph()
        if centered_header:
            jt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = jt.add_run(pi["jobTitle"])
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(90, 90, 90) if centered_header else accent_rgb
        last_header_para = jt

    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location")] if s]
    if contact_parts:
        cp = doc.add_paragraph()
        if centered_header:
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run("  |  ".join(contact_parts) if centered_header else "   ".join(contact_parts))
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(120, 120, 120)
        last_header_para = cp
    if centered_header:
        _docx_bottom_border(last_header_para, accent)

    def add_section(heading: str):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = PtSpace(10)
        r = h.add_run(heading.upper())
        r.font.bold, r.font.name = True, docx_font
        r.font.size = Pt(8.5) if heading_style == "tracked-light" else Pt(10)
        r.font.color.rgb = accent_rgb if heading_style == "tracked-light" else accent_rgb
        if heading_style == "bordered-caps":
            _docx_bottom_border(h, accent)
        return h

    def add_paragraph(text: str, bold: bool = False, grey: bool = False, size: float = 10):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.bold, r.font.size, r.font.name = bold, Pt(size), docx_font
        if grey:
            r.font.color.rgb = RGBColor(120, 120, 120)
        return p

    if "summary" in sections and content.get("summary"):
        if heading_style != "tracked-light":
            add_section("Professional Summary")
        add_paragraph(content["summary"], grey=(heading_style == "tracked-light"))

    if "experience" in sections and content.get("experience"):
        add_section("Experience")
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            date_str = f"  ({start} - {end})" if start else ""
            add_paragraph(position + (f" — {company}" if company else "") + date_str, bold=True)
            for b in exp.get("bullets", []):
                if str(b).strip():
                    add_paragraph(f"  -  {b}")

    if "projects" in sections and content.get("projects"):
        add_section("Projects")
        for proj in content["projects"]:
            name_, tech = proj.get("name", ""), proj.get("technologies", "")
            label = name_ + (f"  ({tech})" if tech else "")
            if label.strip():
                add_paragraph(label, bold=True)
            if proj.get("description"):
                add_paragraph(proj["description"])

    if "education" in sections and content.get("education"):
        add_section("Education")
        for edu in content["education"]:
            inst, end_date = edu.get("institution", ""), edu.get("endDate", "")
            deg = ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            add_paragraph(inst + (f"  ({end_date})" if end_date else ""), bold=True)
            if deg:
                add_paragraph(deg)

    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            add_section("Skills")
            add_paragraph("  -  ".join(skill_names))

    if "certifications" in sections and content.get("certifications"):
        add_section("Certifications")
        for cert in content["certifications"]:
            name_, issuer, date = cert.get("name", ""), cert.get("issuer", ""), cert.get("date", "")
            line = name_ + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                add_paragraph(line)

    if "achievements" in sections and content.get("achievements"):
        add_section("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                add_paragraph(f"  -  {a}")

    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            add_section("Languages")
            add_paragraph("  -  ".join(lang_strs))

    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            add_section("Interests")
            add_paragraph("  -  ".join(interests))

    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title, cs_body = (cs.get("title") or "").strip(), (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            add_section(cs_title or "Additional Information")
            if cs_body:
                add_paragraph(cs_body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _classic_docx_sidebar(content: dict, sections: list[str], *, accent: tuple[int, int, int], docx_font: str,
                           font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Modern: a 1-row, 2-column layout table — left cell shaded with the
    accent color (contact/skills/languages, white text), right cell white
    (summary/experience/projects/education/customSections). Word tables
    auto-paginate, so unlike the PDF version this has no single-page caveat.
    ModernTemplate.tsx doesn't render achievements/certifications/interests;
    like the PDF version, they're still included here rather than silently
    dropping real candidate data. font_scale/spacing_scale come from the
    candidate's own Font-Size/Spacing picker choice (see _resolve_style)."""
    from docx import Document
    from docx.shared import RGBColor, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    Pt = _docx_pt_scaler(font_scale)
    PtSpace = _docx_space_scaler(spacing_scale)

    pi = content.get("personalInfo", {})
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(24)
        sec.left_margin = sec.right_margin = Pt(24)
    doc.styles["Normal"].font.name = docx_font
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.line_spacing = spacing_scale
    accent_rgb = RGBColor(*accent)
    hex_accent = "%02X%02X%02X" % accent

    table = doc.add_table(rows=1, cols=2)
    _docx_no_borders(table)
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.3)
    side_cell, main_cell = table.rows[0].cells
    side_cell.width, main_cell.width = Inches(2.2), Inches(4.3)
    _docx_shade_cell(side_cell, hex_accent)
    side_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def side_para(text: str, size: float = 9, bold: bool = False, color=RGBColor(0xFF, 0xFF, 0xFF)):
        p = side_cell.add_paragraph()
        r = p.add_run(text)
        r.font.size, r.font.bold, r.font.name, r.font.color.rgb = Pt(size), bold, docx_font, color
        return p

    # clear the cell's auto-inserted empty paragraph by reusing it for the name
    side_cell.paragraphs[0].text = ""
    p0 = side_cell.paragraphs[0]
    r0 = p0.add_run(pi.get("fullName") or content.get("title") or "")
    r0.font.size, r0.font.bold, r0.font.name = Pt(14), True, docx_font
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    if pi.get("jobTitle"):
        side_para(pi["jobTitle"], size=9.5, color=RGBColor(0xC8, 0xD2, 0xF0))
    side_para("")

    contact_lines = [v for v in [pi.get("location"), pi.get("phone"), pi.get("email"), pi.get("linkedin"), pi.get("github")] if v]
    if contact_lines:
        side_para("CONTACT", bold=True, size=8.5, color=RGBColor(0xBE, 0xCD, 0xF0))
        for c in contact_lines:
            side_para(c)
        side_para("")

    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            side_para("SKILLS", bold=True, size=8.5, color=RGBColor(0xBE, 0xCD, 0xF0))
            for s in skill_names:
                side_para(f"- {s}")
            side_para("")

    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            side_para("LANGUAGES", bold=True, size=8.5, color=RGBColor(0xBE, 0xCD, 0xF0))
            for l in lang_strs:
                side_para(l)

    main_cell.paragraphs[0].text = ""

    def add_section(heading: str):
        h = main_cell.add_paragraph()
        h.paragraph_format.space_before = PtSpace(8)
        r = h.add_run(heading.upper())
        r.font.bold, r.font.size, r.font.name = True, Pt(10), docx_font
        r.font.color.rgb = accent_rgb
        _docx_bottom_border(h, accent)
        return h

    def add_paragraph(text: str, bold: bool = False, grey: bool = False, size: float = 10):
        p = main_cell.add_paragraph()
        r = p.add_run(text)
        r.font.bold, r.font.size, r.font.name = bold, Pt(size), docx_font
        if grey:
            r.font.color.rgb = RGBColor(120, 120, 120)
        return p

    first = True
    def maybe_use_first(heading):
        nonlocal first
        if first:
            first = False
            h = main_cell.paragraphs[0]
            r = h.add_run(heading.upper())
            r.font.bold, r.font.size, r.font.name = True, Pt(10), docx_font
            r.font.color.rgb = accent_rgb
            _docx_bottom_border(h, accent)
            return h
        return add_section(heading)

    if "summary" in sections and content.get("summary"):
        maybe_use_first("Professional Summary")
        add_paragraph(content["summary"])
    if "experience" in sections and content.get("experience"):
        maybe_use_first("Work Experience")
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            date_str = f"  ({start} - {end})" if start else ""
            add_paragraph(position + (f" - {company}" if company else "") + date_str, bold=True)
            for b in exp.get("bullets", []):
                if str(b).strip():
                    add_paragraph(f"  -  {b}")
    if "projects" in sections and content.get("projects"):
        maybe_use_first("Projects")
        for proj in content["projects"]:
            name_, tech = proj.get("name", ""), proj.get("technologies", "")
            label = name_ + (f"  ({tech})" if tech else "")
            if label.strip():
                add_paragraph(label, bold=True)
            if proj.get("description"):
                add_paragraph(proj["description"])
    if "education" in sections and content.get("education"):
        maybe_use_first("Education")
        for edu in content["education"]:
            inst, end_date = edu.get("institution", ""), edu.get("endDate", "")
            deg = ", ".join(filter(None, [edu.get("degree"), edu.get("field")]))
            add_paragraph(inst + (f"  ({end_date})" if end_date else ""), bold=True)
            if deg:
                add_paragraph(deg)
    if "certifications" in sections and content.get("certifications"):
        maybe_use_first("Certifications")
        for cert in content["certifications"]:
            name_, issuer, date = cert.get("name", ""), cert.get("issuer", ""), cert.get("date", "")
            line = name_ + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                add_paragraph(line)
    if "achievements" in sections and content.get("achievements"):
        maybe_use_first("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                add_paragraph(f"  -  {a}")
    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            maybe_use_first("Interests")
            add_paragraph("  -  ".join(interests))
    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title, cs_body = (cs.get("title") or "").strip(), (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            maybe_use_first(cs_title or "Additional Information")
            if cs_body:
                add_paragraph(cs_body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _classic_docx_sidebar_dark(content: dict, sections: list[str], *, accent: tuple[int, int, int], docx_font: str,
                                font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Executive: a full-width dark header row (its own 1x1 table) then a
    1-row, 2-column layout table below (main column + light-gray sidebar) —
    DOCX counterpart of _classic_pdf_sidebar_dark. font_scale/spacing_scale
    come from the candidate's own Font-Size/Spacing picker choice (see
    _resolve_style)."""
    from docx import Document
    from docx.shared import RGBColor, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL
    Pt = _docx_pt_scaler(font_scale)
    PtSpace = _docx_space_scaler(spacing_scale)

    pi = content.get("personalInfo", {})
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(0)
        sec.left_margin = sec.right_margin = Pt(0)
    doc.styles["Normal"].font.name = docx_font
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.line_spacing = spacing_scale
    accent_rgb = RGBColor(*accent)
    hex_accent = "%02X%02X%02X" % accent

    # Header band
    band = doc.add_table(rows=1, cols=1)
    _docx_no_borders(band)
    band.autofit = False
    band_cell = band.rows[0].cells[0]
    band_cell.width = Inches(8.5)
    _docx_shade_cell(band_cell, hex_accent)
    band_cell.paragraphs[0].paragraph_format.space_before = PtSpace(6)
    band_cell.paragraphs[0].paragraph_format.left_indent = Pt(16)
    r = band_cell.paragraphs[0].add_run((pi.get("fullName") or content.get("title") or "").upper())
    r.font.size, r.font.bold, r.font.name, r.font.color.rgb = Pt(17), True, docx_font, RGBColor(0xFF, 0xFF, 0xFF)
    if pi.get("jobTitle"):
        p = band_cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        r = p.add_run(pi["jobTitle"].upper())
        r.font.size, r.font.name, r.font.color.rgb = Pt(10), docx_font, RGBColor(0xC8, 0xC8, 0xCD)
    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location"), pi.get("linkedin")] if s]
    if contact_parts:
        p = band_cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.space_after = PtSpace(8)
        r = p.add_run("    ".join(contact_parts))
        r.font.size, r.font.name, r.font.color.rgb = Pt(8.5), docx_font, RGBColor(0xB4, 0xB4, 0xBA)

    doc.add_paragraph().paragraph_format.space_after = PtSpace(4)

    grid = doc.add_table(rows=1, cols=2)
    _docx_no_borders(grid)
    grid.autofit = False
    grid.columns[0].width = Inches(5.3)
    grid.columns[1].width = Inches(3.2)
    main_cell, side_cell = grid.rows[0].cells
    main_cell.width, side_cell.width = Inches(5.3), Inches(3.2)
    _docx_shade_cell(side_cell, "F8F9FA")
    main_cell.paragraphs[0].paragraph_format.left_indent = Pt(16)
    side_cell.paragraphs[0].paragraph_format.left_indent = Pt(10)

    def divider_heading(cell, text: str):
        h = cell.add_paragraph() if cell.paragraphs[0].runs else cell.paragraphs[0]
        h.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
        h.paragraph_format.space_before = PtSpace(8)
        r = h.add_run(text.upper())
        r.font.bold, r.font.size, r.font.name = True, Pt(8.5), docx_font
        r.font.color.rgb = accent_rgb
        _docx_bottom_border(h, (210, 210, 214))
        return h

    def side_heading(cell, text: str):
        h = cell.add_paragraph() if cell.paragraphs[0].runs else cell.paragraphs[0]
        h.paragraph_format.left_indent = Pt(10)
        h.paragraph_format.space_before = PtSpace(8)
        r = h.add_run(text.upper())
        r.font.bold, r.font.size, r.font.name = True, Pt(8), docx_font
        r.font.color.rgb = RGBColor(0x5A, 0x5A, 0x5A)
        return h

    def cell_body(cell, text: str, bold: bool = False, grey: bool = False, size: float = 9.5, indent: float = 16):
        p = cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(indent)
        r = p.add_run(text)
        r.font.bold, r.font.size, r.font.name = bold, Pt(size), docx_font
        if grey:
            r.font.color.rgb = RGBColor(120, 120, 120)
        return p

    if "summary" in sections and content.get("summary"):
        divider_heading(main_cell, "Executive Profile")
        cell_body(main_cell, content["summary"])
    if "experience" in sections and content.get("experience"):
        divider_heading(main_cell, "Career History")
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            cell_body(main_cell, position, bold=True)
            meta = company + (f"  ({start} - {end})" if start else "")
            if meta.strip():
                cell_body(main_cell, meta, grey=True, size=8.5)
            for b in exp.get("bullets", []):
                if str(b).strip():
                    cell_body(main_cell, f"  -  {b}")
    if "projects" in sections and content.get("projects"):
        divider_heading(main_cell, "Projects")
        for proj in content["projects"]:
            name_, tech = proj.get("name", ""), proj.get("technologies", "")
            label = name_ + (f"  ({tech})" if tech else "")
            if label.strip():
                cell_body(main_cell, label, bold=True)
            if proj.get("description"):
                cell_body(main_cell, proj["description"])
    if "achievements" in sections and content.get("achievements"):
        divider_heading(main_cell, "Key Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                cell_body(main_cell, f"  -  {a}")
    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title, cs_body = (cs.get("title") or "").strip(), (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            divider_heading(main_cell, cs_title or "Additional Information")
            if cs_body:
                cell_body(main_cell, cs_body)

    if "education" in sections and content.get("education"):
        side_heading(side_cell, "Education")
        for edu in content["education"]:
            cell_body(side_cell, edu.get("degree", ""), bold=True, size=9, indent=10)
            if edu.get("field"):
                cell_body(side_cell, edu["field"], grey=True, size=8.5, indent=10)
            if edu.get("institution"):
                cell_body(side_cell, edu["institution"], grey=True, size=8.5, indent=10)
    if "skills" in sections:
        skill_names = _skill_names(content)
        if skill_names:
            side_heading(side_cell, "Core Competencies")
            for s in skill_names:
                cell_body(side_cell, f"- {s}", size=8.5, indent=10)
    if "certifications" in sections and content.get("certifications"):
        side_heading(side_cell, "Certifications")
        for cert in content["certifications"]:
            cell_body(side_cell, cert.get("name", ""), bold=True, size=8.5, indent=10)
            meta = " · ".join(filter(None, [cert.get("issuer"), cert.get("date")]))
            if meta:
                cell_body(side_cell, meta, grey=True, size=8, indent=10)
    # Not rendered by ExecutiveTemplate.tsx's Preview at all -- appended here
    # so real candidate data is never silently dropped (see Modern's docstring).
    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            side_heading(side_cell, "Languages")
            for l in lang_strs:
                cell_body(side_cell, l, size=8.5, indent=10)
    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            side_heading(side_cell, "Interests")
            cell_body(side_cell, "  -  ".join(interests), size=8.5, indent=10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _classic_docx_header_band(content: dict, sections: list[str], *, accent: tuple[int, int, int], docx_font: str,
                               font_scale: float = 1.0, spacing_scale: float = 1.0) -> bytes:
    """Creative: a solid-accent header band (DOCX can't do the frontend's
    purple→blue gradient — documented simplification), then single-column
    content with accent-colored bold sub-text standing in for the frontend's
    left-border-accented entries (a real per-paragraph left border is
    possible in DOCX via oxml but adds meaningfully more complexity for a
    cosmetic-only difference), and a 2-column Education | Skills table near
    the end. font_scale/spacing_scale come from the candidate's own
    Font-Size/Spacing picker choice (see _resolve_style)."""
    from docx import Document
    from docx.shared import RGBColor, Inches
    Pt = _docx_pt_scaler(font_scale)
    PtSpace = _docx_space_scaler(spacing_scale)

    pi = content.get("personalInfo", {})
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Pt(0)
        sec.left_margin = sec.right_margin = Pt(0)
    doc.styles["Normal"].font.name = docx_font
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.line_spacing = spacing_scale
    accent_rgb = RGBColor(*accent)
    hex_accent = "%02X%02X%02X" % accent

    band = doc.add_table(rows=1, cols=1)
    _docx_no_borders(band)
    band.autofit = False
    band_cell = band.rows[0].cells[0]
    band_cell.width = Inches(8.5)
    _docx_shade_cell(band_cell, hex_accent)
    band_cell.paragraphs[0].paragraph_format.space_before = PtSpace(6)
    band_cell.paragraphs[0].paragraph_format.left_indent = Pt(16)
    r = band_cell.paragraphs[0].add_run(pi.get("fullName") or content.get("title") or "")
    r.font.size, r.font.bold, r.font.name, r.font.color.rgb = Pt(16), True, docx_font, RGBColor(0xFF, 0xFF, 0xFF)
    if pi.get("jobTitle"):
        p = band_cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        r = p.add_run(pi["jobTitle"])
        r.font.size, r.font.name, r.font.color.rgb = Pt(10), docx_font, RGBColor(0xE6, 0xE1, 0xFA)
    contact_parts = [s for s in [pi.get("email"), pi.get("phone"), pi.get("location")] if s]
    if contact_parts:
        p = band_cell.add_paragraph()
        p.paragraph_format.left_indent = Pt(16)
        p.paragraph_format.space_after = PtSpace(8)
        r = p.add_run("    ".join(contact_parts))
        r.font.size, r.font.name, r.font.color.rgb = Pt(8.5), docx_font, RGBColor(0xE1, 0xDC, 0xFA)

    body_indent = Pt(16)

    def add_section(heading: str):
        h = doc.add_paragraph()
        h.paragraph_format.left_indent = body_indent
        h.paragraph_format.space_before = PtSpace(8)
        r = h.add_run(heading.upper())
        r.font.bold, r.font.size, r.font.name, r.font.color.rgb = True, Pt(8.5), docx_font, accent_rgb
        return h

    def add_paragraph(text: str, bold: bool = False, accent_text: bool = False, size: float = 9.5):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = body_indent
        r = p.add_run(text)
        r.font.bold, r.font.size, r.font.name = bold, Pt(size), docx_font
        if accent_text:
            r.font.color.rgb = accent_rgb
        return p

    if "summary" in sections and content.get("summary"):
        add_section("About Me")
        add_paragraph(content["summary"])
    if "experience" in sections and content.get("experience"):
        add_section("Experience")
        for exp in content["experience"]:
            position, company = exp.get("position", ""), exp.get("company", "")
            start = exp.get("startDate", "")
            end = "Present" if exp.get("current") else exp.get("endDate", "")
            add_paragraph(position, bold=True)
            meta = company + (f"  ({start} - {end})" if start else "")
            if meta.strip():
                add_paragraph(meta, accent_text=True)
            for b in exp.get("bullets", []):
                if str(b).strip():
                    add_paragraph(f"›  {b}")
    if "projects" in sections and content.get("projects"):
        add_section("Projects")
        for proj in content["projects"]:
            name_, tech = proj.get("name", ""), proj.get("technologies", "")
            add_paragraph(name_, bold=True)
            if tech:
                add_paragraph(tech, accent_text=True)
            if proj.get("description"):
                add_paragraph(proj["description"])

    has_edu = "education" in sections and content.get("education")
    skill_names = _skill_names(content) if "skills" in sections else []
    if has_edu or skill_names:
        grid = doc.add_table(rows=1, cols=2)
        _docx_no_borders(grid)
        grid.autofit = False
        grid.columns[0].width = Inches(4.0)
        grid.columns[1].width = Inches(4.0)
        edu_cell, skill_cell = grid.rows[0].cells
        edu_cell.width, skill_cell.width = Inches(4.0), Inches(4.0)
        edu_cell.paragraphs[0].paragraph_format.left_indent = body_indent
        skill_cell.paragraphs[0].paragraph_format.left_indent = Pt(8)
        if has_edu:
            h = edu_cell.paragraphs[0]
            r = h.add_run("EDUCATION")
            r.font.bold, r.font.size, r.font.name, r.font.color.rgb = True, Pt(8.5), docx_font, accent_rgb
            for edu in content["education"]:
                p = edu_cell.add_paragraph()
                p.paragraph_format.left_indent = body_indent
                r = p.add_run(edu.get("degree", ""))
                r.font.bold, r.font.size, r.font.name = True, Pt(9), docx_font
                if edu.get("institution"):
                    p2 = edu_cell.add_paragraph()
                    p2.paragraph_format.left_indent = body_indent
                    r2 = p2.add_run(edu["institution"])
                    r2.font.size, r2.font.name, r2.font.color.rgb = Pt(8.5), docx_font, RGBColor(120, 120, 120)
        if skill_names:
            h = skill_cell.paragraphs[0]
            r = h.add_run("SKILLS")
            r.font.bold, r.font.size, r.font.name, r.font.color.rgb = True, Pt(8.5), docx_font, accent_rgb
            for s in skill_names:
                p = skill_cell.add_paragraph()
                p.paragraph_format.left_indent = Pt(8)
                r = p.add_run(s)
                r.font.size, r.font.name = Pt(8.5), docx_font

    if "certifications" in sections and content.get("certifications"):
        add_section("Certifications")
        for cert in content["certifications"]:
            name_, issuer, date = cert.get("name", ""), cert.get("issuer", ""), cert.get("date", "")
            line = name_ + (f" — {issuer}" if issuer else "") + (f" ({date})" if date else "")
            if line.strip():
                add_paragraph(line)
    if "achievements" in sections and content.get("achievements"):
        add_section("Achievements")
        for a in content["achievements"]:
            if str(a).strip():
                add_paragraph(f"›  {a}")
    if "languages" in sections and content.get("languages"):
        lang_strs = _language_strings(content)
        if lang_strs:
            add_section("Languages")
            add_paragraph("  -  ".join(lang_strs))
    if "interests" in sections and content.get("interests"):
        interests = [str(i) for i in content["interests"] if str(i).strip()]
        if interests:
            add_section("Interests")
            add_paragraph("  -  ".join(interests))
    if "customSections" in sections and content.get("customSections"):
        for cs in content["customSections"]:
            cs_title, cs_body = (cs.get("title") or "").strip(), (cs.get("content") or "").strip()
            if not cs_title and not cs_body:
                continue
            add_section(cs_title or "Additional Information")
            if cs_body:
                add_paragraph(cs_body)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_classic(content: dict, title: str, sections: list[str], *, template_id: str,
                         style: dict | None = None) -> bytes:
    """DOCX counterpart of _build_pdf_classic — same dispatch-by-layoutFamily
    contract, same TEMPLATE_SPECS single source of truth, same `style`
    override contract (see _resolve_style)."""
    spec = TEMPLATE_SPECS.get(template_id, {})
    style = style or {}
    accent = _hex_to_rgb(spec.get("accent") or "#1e40af")
    font_family = style.get("font_family") or spec.get("font") or "sans"
    docx_font = style.get("docx_font") or _DOCX_FONT_MAP.get(font_family, "Calibri")
    font_scale = style.get("font_scale", 1.0)
    spacing_scale = style.get("spacing_scale", 1.0)
    layout = spec.get("layoutFamily") or ""

    if layout == "sidebar-left":
        return _classic_docx_sidebar(content, sections, accent=accent, docx_font=docx_font,
                                      font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "sidebar-dark":
        return _classic_docx_sidebar_dark(content, sections, accent=accent, docx_font=docx_font,
                                           font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "header-band":
        return _classic_docx_header_band(content, sections, accent=accent, docx_font=docx_font,
                                          font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "single-column-serif":
        return _classic_docx_single_column(content, sections, accent=accent, docx_font=docx_font,
                                            centered_header=True, heading_style="bordered-caps",
                                            font_scale=font_scale, spacing_scale=spacing_scale)
    if layout == "single-column-minimal":
        return _classic_docx_single_column(content, sections, accent=accent, docx_font=docx_font,
                                            centered_header=False, heading_style="tracked-light",
                                            font_scale=font_scale, spacing_scale=spacing_scale)
    return _build_docx(content, title, sections)


def _make_classic_builder(template_id: str) -> "TemplateBuilder":
    def pdf(content: dict, title: str, sections: list[str], style: dict | None = None) -> bytes:
        return _build_pdf_classic(content, title, sections, template_id=template_id, style=style)

    def docx(content: dict, title: str, sections: list[str], style: dict | None = None) -> bytes:
        return _build_docx_classic(content, title, sections, template_id=template_id, style=style)

    return TemplateBuilder(pdf=pdf, docx=docx, layout_family=TEMPLATE_SPECS[template_id]["layoutFamily"])


def _make_single_column_builder(template_id: str) -> "TemplateBuilder":
    def pdf(content: dict, title: str, sections: list[str], style: dict | None = None) -> bytes:
        return _build_pdf_single_column(content, title, sections, template_id=template_id, style=style)

    def docx(content: dict, title: str, sections: list[str], style: dict | None = None) -> bytes:
        return _build_docx_single_column(content, title, sections, template_id=template_id, style=style)

    return TemplateBuilder(pdf=pdf, docx=docx, layout_family="single-column")


# ── Template registry ───────────────────────────────────────────────────────
# Every template_id now gets a builder whose PDF/DOCX output actually reflects
# its own accent/font/layoutFamily (shared/template-specs.json — the single
# configuration layer read by both the frontend template picker and this
# exporter) PLUS the candidate's own explicit font/spacing override (`style`
# — see _resolve_style; the SAME Resume.font_metadata/layout_metadata the
# Resume Editor's Font/Spacing picker writes): modern/professional/minimal/
# creative/executive route through _make_classic_builder (five genuinely
# distinct "classic-family" layouts — see _build_pdf_classic/
# _build_docx_classic), the five Phase 2 templates through
# _make_single_column_builder (SINGLE_COLUMN_CONFIGS above). Adding a
# genuinely new layout shape later means adding a new family function and
# repointing that id's entry here; callers (export_pdf/export_docx,
# routers/resumes.py::download_resume) never change. `style` defaults to
# None (== today's unscaled, template-default-font behavior) everywhere, so
# every existing direct caller (tests, other call sites) that doesn't pass
# it keeps working unchanged.

@dataclass
class TemplateBuilder:
    pdf: Callable[[dict, str, list[str], dict | None], bytes]
    docx: Callable[[dict, str, list[str], dict | None], bytes]
    layout_family: str


TEMPLATE_BUILDERS: dict[str, TemplateBuilder] = {
    template_id: (
        _make_single_column_builder(template_id) if template_id in SINGLE_COLUMN_CONFIGS
        else _make_classic_builder(template_id)
    )
    for template_id, spec in TEMPLATE_SPECS.items()
}
